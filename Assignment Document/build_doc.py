# -*- coding: utf-8 -*-
"""Builds the AI Smart Inventory Management System project report (.docx)
following the structure of the college's MERN_PROJECT_DOCUMENTATION template."""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Users\Admin\Desktop\Clg Mern\Assignment Document"
DIAG = os.path.join(BASE, "diagrams")
OUT  = os.environ.get("OUT_DOCX", os.path.join(BASE, "AI-Smart-Inventory-Documentation.docx"))

NAVY  = RGBColor(0x1F, 0x38, 0x64)
GREY  = RGBColor(0x40, 0x40, 0x40)

doc = Document()

# ---------------------------------------------------------------- page setup
sec = doc.sections[0]
sec.page_height = Inches(11.69)   # A4
sec.page_width  = Inches(8.27)
sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1.1); sec.right_margin = Inches(1)

def set_style_font(style, font="Times New Roman"):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), font)

# Normal
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'; normal.font.size = Pt(12)
normal.font.color.rgb = RGBColor(0, 0, 0)
set_style_font(normal)
pf = normal.paragraph_format
pf.line_spacing = 1.5; pf.space_after = Pt(6); pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Headings
for name, size in [('Heading 1', 16), ('Heading 2', 13.5), ('Heading 3', 12.5)]:
    st = doc.styles[name]
    st.font.name = 'Times New Roman'; st.font.size = Pt(size)
    st.font.bold = True; st.font.color.rgb = NAVY
    set_style_font(st)
    st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.keep_with_next = True
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ---------------------------------------------------------------- helpers
def body(text, align='justify', bold=False, italic=False, size=12, color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT, 'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.4 + 0.3 * level)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # support a leading bold lead-in "Term: rest"
    if '||' in text:
        lead, rest = text.split('||', 1)
        r1 = p.add_run(lead); r1.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p

_num_counter = [0]
def reset_num():
    _num_counter[0] = 0

def numbered(text):
    # Manual numbering with a hanging indent so each list restarts at 1
    # (avoids Word's List Number style continuing across separate lists).
    _num_counter[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("%d.  " % _num_counter[0])
    if '||' in text:
        lead, rest = text.split('||', 1)
        r1 = p.add_run(lead); r1.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p

def h1(text):
    return doc.add_paragraph(text, style='Heading 1')

def h2(text):
    return doc.add_paragraph(text, style='Heading 2')

def h3(text):
    return doc.add_paragraph(text, style='Heading 3')

def page_break():
    doc.add_page_break()

def chapter_banner(num):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("CHAPTER %d" % num); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = GREY

# Master list of figures, in the exact order they appear in the document.
# figure() auto-numbers from this list and asserts the title matches.
FIG_TITLES = [
    "System Module Diagram",                       # 1
    "Core Business Data Flow",                     # 2
    "Three-Tier Architecture of the System",       # 3
    "The MERN Stack",                              # 4
    "MERN Application Hierarchy",                   # 5
    "Technologies Used in the Project",            # 6
    "A Simple React Component",                     # 7
    "A React Hook Example",                         # 8
    "A React Context Provider",                     # 9
    "Context API State Management",                 # 10
    "A Simple Node.js Web Server",                  # 11
    "An Express Server",                            # 12
    "MongoDB Atlas Cluster",                        # 13
    "Client-Side Dependencies (package.json)",     # 14
    "Server-Side Dependencies (package.json)",     # 15
    "JWT Authentication Flow",                      # 16
    "Database Schema and Relationships",           # 17
]
_fig_counter = [0]

def figure(img, title, width=6.2):
    _fig_counter[0] += 1
    n = _fig_counter[0]
    assert title == FIG_TITLES[n - 1], \
        "Figure %d title mismatch: %r vs expected %r" % (n, title, FIG_TITLES[n - 1])
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(os.path.join(DIAG, img), width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    r = c.add_run("Figure %d — %s" % (n, title))
    r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = GREY

def centered(text, bold=False, size=12, color=None, space_after=4, space_before=0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    return p

def front_title(text):
    """Centered front-matter section title (NOT a heading -> not in TOC)."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
    return p

def add_page_number_footer():
    footer = sec.footer
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)
    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = 'TOC \\o "1-3" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "Update this field in Word (select it and press F9) to generate the Table of Contents."
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2); run._r.append(t); run._r.append(f3)

add_page_number_footer()

# ================================================================ COVER PAGE
for _ in range(1): doc.add_paragraph()
centered("AI SMART INVENTORY MANAGEMENT SYSTEM", bold=True, size=24, color=NAVY, space_after=2)
centered("USING MERN STACK", bold=True, size=18, color=NAVY, space_after=14)
centered("A Project Report Submitted", size=12.5, space_after=2)
centered("in Partial Fulfillment of the Requirements for the Subject", size=12.5, space_after=2)
centered("PROJECT", bold=True, size=12.5, space_after=2)
centered("of", size=12.5, space_after=2)
centered("MASTER OF COMPUTER APPLICATIONS (MCA) – II", bold=True, size=13, space_after=16)
centered("SUBMITTED BY:", bold=True, size=12.5, space_after=2)
centered("[ STUDENT NAME ]", size=12.5, space_after=2)
centered("[ ENROLLMENT NUMBER ]", size=12.5, space_after=16)
centered("Under the Guidance of", size=12.5, space_after=2)
centered("PROF. HIMANI KHODIFAD", bold=True, size=13, space_after=18)
centered("R.B. Institute of Management Studies (RBIMS)", bold=True, size=12.5, space_after=2)
centered("APPROVED BY AICTE, GOVT. OF INDIA, NEW DELHI", size=11, space_after=2)
centered("Affiliated with Gujarat Technological University, Ahmedabad", size=11, space_after=2)
centered("Ahmedabad, Gujarat", size=11, space_after=10)
centered("Academic Year: 2025–26", bold=True, size=12.5)
page_break()

# ================================================================ DECLARATION
front_title("DECLARATION")
body("I hereby declare that the work presented in this project report entitled “AI Smart Inventory "
     "Management System using MERN Stack” submitted in partial fulfillment of the requirements for the "
     "subject Project of MCA Semester – II under Gujarat Technological University (GTU) is an authentic "
     "record of the work carried out by me during the academic year 2025–26 under the guidance of "
     "Prof. Himani Khodifad.")
body("The matter embodied in this report has not been submitted by me for the award of any other degree, "
     "diploma, or similar title of this or any other University/Institute.")
doc.add_paragraph(); doc.add_paragraph()
body("[ STUDENT NAME ]", align='right', bold=True, space_after=2)
body("[ ENROLLMENT NUMBER ]", align='right', bold=True)
doc.add_paragraph()
body("This is to certify that the above statement made by the student is correct to the best of our knowledge.", italic=True)
doc.add_paragraph()
body("Project Guide", bold=True, space_after=2)
body("Prof. Himani Khodifad", space_after=2)
body("Signature: __________________", space_after=2)
body("Date: __________________", space_after=2)
body("Place: RBIMS, Ahmedabad", space_after=2)
page_break()

# ================================================================ CERTIFICATE
centered("R.B. Institute of Management Studies (RBIMS)", bold=True, size=13, color=NAVY, space_after=2)
centered("APPROVED BY AICTE, GOVT. OF INDIA, NEW DELHI", size=10.5, space_after=2)
centered("Affiliated with Gujarat Technological University, Ahmedabad", size=10.5, space_after=14)
front_title("CERTIFICATE")
body("This is to certify that the project entitled “AI Smart Inventory Management System using MERN "
     "Stack” submitted by the following student of MCA Semester – II of R.B. Institute of Management "
     "Studies (RBIMS), affiliated with Gujarat Technological University (GTU), Ahmedabad, is a bonafide "
     "record of original project work carried out under the guidance and supervision of Prof. Himani "
     "Khodifad during the academic year 2025–26.")
doc.add_paragraph()
body("[ STUDENT NAME ]", bold=True, space_after=2)
body("[ ENROLLMENT NUMBER ]", bold=True)
doc.add_paragraph()
body("This work has not been submitted elsewhere for the award of any degree, diploma, or similar title.")
doc.add_paragraph(); doc.add_paragraph()
body("Project Guide", bold=True, space_after=2)
body("Prof. Himani Khodifad", space_after=2)
body("Signature: __________________")
doc.add_paragraph()
body("Head of Department", bold=True, space_after=2)
body("Signature: __________________")
doc.add_paragraph()
body("R.B. Institute of Management Studies (RBIMS), Ahmedabad, Gujarat", space_after=2)
body("Date: __________________", space_after=2)
page_break()

# ================================================================ ABSTRACT
front_title("ABSTRACT")
body("The AI Smart Inventory Management System is a full-stack web application designed to digitise the "
     "day-to-day operations of small and medium retail businesses such as grocery (kirana) stores, traders "
     "and wholesalers. A large proportion of such businesses continue to manage stock in paper registers, "
     "prepare bills manually, and record customer credit in physical ledgers. These practices result in "
     "calculation errors, lost records, poor tax compliance and an absence of meaningful business insight.")
body("The proposed system addresses these problems through a single integrated platform built on the MERN "
     "stack — MongoDB, Express.js, React and Node.js. It provides real-time inventory tracking, rapid "
     "point-of-sale billing with automatic GST computation, server-generated PDF invoices and a UPI "
     "“Scan-to-Pay” QR code, a digital customer credit ledger (khata), and supplier management. The "
     "application also provides intelligent assistance: demand forecasting and reorder suggestions are "
     "computed from sales history, an AI help assistant powered by the Google Gemini API answers "
     "natural-language questions about the business, and supplier bills are read automatically using "
     "Gemini Vision with a Tesseract.js optical-character-recognition fallback. Scheduled background jobs "
     "generate periodic reports and raise smart stock alerts automatically.")
body("The system follows a layered three-tier architecture with validated REST APIs, JWT-based "
     "authentication, and precise decimal arithmetic for all monetary calculations. The result is a "
     "reliable, extensible and practical system that demonstrates the complete software development life "
     "cycle — from requirement analysis and database design to implementation and testing.")
page_break()

# ================================================================ ACKNOWLEDGEMENT
front_title("ACKNOWLEDGEMENT")
body("The completion of this project has been possible through the support and encouragement of many "
     "individuals, and I take this opportunity to express my sincere gratitude to all of them.")
body("First and foremost, I express my heartfelt thanks to my project guide, Prof. Himani Khodifad, for "
     "her invaluable guidance, constant encouragement and constructive feedback throughout the course of "
     "this project. Her insights and suggestions were instrumental in shaping this work.")
body("I am thankful to the Head of the Department and the entire faculty of the Master of Computer "
     "Applications programme at R.B. Institute of Management Studies (RBIMS), Ahmedabad, for providing the "
     "academic environment and resources required to carry out this project.")
body("I also extend my gratitude to Gujarat Technological University (GTU) for including a practical "
     "project component in the curriculum, which gave me the opportunity to apply theoretical knowledge to "
     "a real-world problem.")
body("Finally, I thank my family and friends for their continual support and motivation during the "
     "development of this project.")
doc.add_paragraph(); doc.add_paragraph()
body("[ STUDENT NAME ]", align='right', bold=True)
page_break()

# ================================================================ TABLE OF CONTENTS
front_title("TABLE OF CONTENTS")
add_toc()
page_break()

# ================================================================ LIST OF FIGURES
front_title("LIST OF FIGURES")
for i, t in enumerate(FIG_TITLES, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.2)
    rn = p.add_run("Figure %d" % i); rn.bold = True
    p.add_run("     " + t)
page_break()

# ================================================================ CHAPTER 1
chapter_banner(1)
h1("1. Introduction")
body("Inventory management is the backbone of every retail business. While large enterprises employ "
     "sophisticated Enterprise Resource Planning (ERP) systems, the vast majority of small retailers — "
     "grocery stores, traders and wholesalers — continue to depend on manual methods. Stock registers "
     "are updated by hand, bills are calculated mentally or on a calculator, and customer credit is "
     "maintained in a paper ledger commonly known as a khata.")
body("This project, the AI Smart Inventory Management System, was developed to bring the benefits of modern "
     "software and artificial intelligence to such businesses in a form that is affordable, simple and "
     "aligned with their actual daily workflow. The system unifies inventory, billing, customer credit, "
     "supplier records, analytics and reporting into one web application, and applies AI where it produces "
     "direct business value: predicting demand, recommending reorders and automating routine analysis.")

h2("1.1 Problem Statement")
body("The existing manual approach followed by most small retail businesses suffers from several "
     "interrelated problems:")
bullet("Manual stock-keeping:|| stock levels recorded on paper are frequently inaccurate, and shortages "
       "are discovered only when a customer asks for an item that is unavailable.")
bullet("Error-prone billing:|| manual GST computation leads to incorrect tax amounts, non-compliant "
       "invoices and revenue leakage.")
bullet("Unreliable credit records:|| paper credit ledgers are easily lost or disputed, and outstanding "
       "amounts are difficult to total or follow up.")
bullet("No business insight:|| the owner has no data on fast-moving products, seasonal demand, stock "
       "shrinkage or profitability.")
bullet("Unsuitable existing software:|| commercial solutions are often expensive, overly complex, or not "
       "designed around the workflow of a small Indian retail business.")

h2("1.2 Innovative Ideas of the Project")
body("The project introduces several ideas that distinguish it from a conventional inventory application:")
bullet("Demand forecasting from sales history|| that analyses past sales to predict future demand and "
       "recommend what to reorder and when.")
bullet("An AI help assistant|| (powered by the Google Gemini API) that answers natural-language "
       "questions about the business, with a built-in rule-based fallback when the AI service is "
       "unavailable.")
bullet("Automatic bill reading,|| allowing the owner to photograph a supplier bill and have products "
       "extracted automatically (using Gemini Vision, with a Tesseract.js OCR fallback) instead of "
       "typing them by hand.")
bullet("A fully digital khata|| that replaces the paper credit ledger and keeps an accurate, permanent, "
       "timestamped record of every credit and repayment.")
bullet("Amount-first billing,|| where the customer asks for a fixed value of a product (for example "
       "“Rs. 50 of sugar”) and the system back-calculates the quantity.")
bullet("A UPI “Scan-to-Pay” QR code on the invoice,|| generated from the business’s UPI ID, enabling "
       "the customer to scan and pay instantly.")
bullet("Automation through scheduled jobs|| that raise stock alerts and generate reports without any "
       "manual effort.")

h2("1.3 Project Objective")
body("The principal objectives of the system are:")
reset_num()
numbered("To digitise inventory with real-time stock levels, units of measure, low-stock detection and "
         "shrinkage (stock-adjustment) tracking.")
numbered("To provide a fast point-of-sale billing module with automatic GST calculation, PDF invoice "
         "generation and a UPI Scan-to-Pay QR code for digital payment.")
numbered("To replace the paper credit ledger with a digital khata that maintains an accurate running "
         "balance for every customer.")
numbered("To provide intelligent assistance — demand forecasting and reorder suggestions derived from "
         "sales history, and an AI-powered help assistant built on the Google Gemini API.")
numbered("To automate routine work — periodic reports and stock alerts — using scheduled "
         "background jobs.")
numbered("To deliver analytics dashboards and exportable reports, including a Tally-compatible export "
         "for accountants.")
numbered("To ensure correctness and security through input validation, authenticated APIs and precise "
         "decimal arithmetic for money.")

h2("1.4 Scope of the Project")
body("The system is a web-based application intended for the owner and staff of a small or medium retail "
     "business. Its scope covers the complete operational cycle of such a business: adding stock, selling "
     "and billing, collecting payment or credit, and analysing performance.")
body("The application supports product and inventory management, point-of-sale billing with GST, customer "
     "and supplier management, a digital credit ledger, AI insights, OCR bill scanning, analytics and "
     "automated reporting. It is multi-lingual and designed to be usable by non-technical shopkeepers. "
     "The current scope is limited to a single business account per login and operation over an internet "
     "connection; multi-branch operation, a dedicated mobile application and offline synchronisation are "
     "identified as future work.")

h2("1.5 Design and Implementation")
body("The application is implemented on the MERN stack and follows a clean separation between the "
     "presentation, application and data layers. On the server, functionality is divided into independent "
     "modules, each exposed as a group of REST endpoints and backed by a controller and a service that "
     "contains the business logic. On the client, each module corresponds to one or more React pages that "
     "consume those endpoints. The principal modules of the system are shown in Figure 1.")
figure("d2_modules.png", "System Module Diagram", width=6.4)
body("Every business operation updates all related data in a single coherent workflow. When a sale is "
     "completed the stock is reduced, the invoice is generated, the payment is recorded and — if the "
     "sale is on credit — the customer’s ledger is updated. The accumulated data subsequently "
     "feeds the analytics and AI layer. This core data flow is illustrated in Figure 2.")
figure("d4_dataflow.png", "Core Business Data Flow", width=6.4)

h2("1.6 Three-Tier Architecture")
body("The system is organised into a classic three-tier architecture, which separates the user interface, "
     "the business logic and the data storage into independent layers. This separation improves "
     "maintainability, allows each layer to be developed and tested in isolation, and makes the system "
     "easier to scale and extend.")
bullet("Presentation tier:|| a React single-page application that runs in the browser and renders the "
       "user interface.")
bullet("Application tier:|| a Node.js and Express REST API that authenticates requests, validates input "
       "and executes the business logic through a controller-and-service structure.")
bullet("Data tier:|| a MongoDB database accessed through the Mongoose object-document mapper, which stores "
       "all persistent data.")
figure("d1_architecture.png", "Three-Tier Architecture of the System", width=6.4)
body("On the server, every request passes through a defined pipeline: authentication middleware verifies "
     "the JSON Web Token, validation middleware checks the request body against a Zod schema, the "
     "controller orchestrates the operation, and the service layer performs the business logic and "
     "database access. A central error-handling middleware converts all failures into consistent JSON "
     "error responses.")
page_break()

# ================================================================ CHAPTER 2
chapter_banner(2)
h1("2. Introduction to Full-Stack Development with JavaScript")
body("Full-stack development refers to the practice of building both the client side (front end) and the "
     "server side (back end) of a web application, together with the database that stores its data. A "
     "developer or team that handles all of these layers is said to work across the full stack.")
body("Traditionally, the front end and back end of an application were written in different programming "
     "languages — for example JavaScript in the browser and PHP, Java or Python on the server. The "
     "arrival of Node.js made it possible to run JavaScript on the server as well, which means a single "
     "language can be used across the entire application. This is the foundation of the MERN stack used in "
     "this project.")
body("MERN is an acronym for four JavaScript-based technologies that together cover every layer of a "
     "modern web application:")
bullet("MongoDB|| — a NoSQL document database that stores data as flexible JSON-like documents.")
bullet("Express.js|| — a minimal and flexible web-application framework for Node.js that is used to "
       "build the REST API.")
bullet("React|| — a component-based JavaScript library for building interactive user interfaces.")
bullet("Node.js|| — a JavaScript runtime that executes the server-side code.")
body("These four technologies and their respective roles are summarised in Figure 4.")
figure("c1_mern_overview.png", "The MERN Stack", width=6.3)
body("Using JavaScript end-to-end offers several advantages: a single language reduces the learning curve "
     "and context-switching; data is exchanged between client and server as JSON, which maps naturally to "
     "JavaScript objects and to MongoDB documents; and a large ecosystem of open-source packages is "
     "available through the Node Package Manager (npm). These factors make the MERN stack a popular and "
     "productive choice for building data-driven applications such as the one developed in this project.")
body("The way these layers are arranged, from the browser down to the database, is shown in Figure 5.")
figure("c2_mern_hierarchy.png", "MERN Application Hierarchy", width=5.6)
page_break()

# ================================================================ CHAPTER 3
chapter_banner(3)
h1("3. Technologies and Concepts")
body("This chapter describes the technologies, libraries and concepts used to build the AI Smart Inventory "
     "Management System, and explains the role each one plays in the application.")

body("The principal technologies used in the project, grouped by the layer in which they operate, are "
     "shown in Figure 6.")
figure("c3_tech_used.png", "Technologies Used in the Project", width=6.2)

h2("3.1 React JS")
body("React is an open-source JavaScript library, developed by Meta (Facebook), for building user "
     "interfaces. It follows a component-based model in which the interface is broken down into small, "
     "reusable pieces called components, each managing its own structure and behaviour. React uses a "
     "virtual DOM to update only the parts of the page that actually change, which makes the interface "
     "fast and responsive. In this project, React (version 19) together with the Vite build tool and the "
     "Tailwind CSS framework is used to build the entire front end — the dashboard, inventory, billing, "
     "khata, analytics and other pages.")

body("Figure 7 shows a simple React functional component that displays the details of a single product.")
figure("c4_react_component.png", "A Simple React Component", width=5.7)

h2("3.2 React Hooks")
body("Hooks are functions introduced in React that let functional components use state and other React "
     "features without writing class components. The most commonly used hooks are useState, which adds "
     "local state to a component, and useEffect, which performs side effects such as fetching data when a "
     "component renders. This project also uses useContext to consume shared application state and a number "
     "of custom hooks (for example, hooks for inventory, suppliers and transactions) that encapsulate data "
     "fetching and reuse it across pages.")

body("Figure 8 shows a custom hook that uses the useState and useEffect hooks to load data from the API.")
figure("c5_react_hook.png", "A React Hook Example", width=5.7)

h2("3.3 State Management — React Context API")
body("As an application grows, data such as the logged-in user, the selected theme and workspace settings "
     "need to be shared across many components. While libraries such as Redux are often used for this "
     "purpose in large applications, this project uses React’s built-in Context API, which provides a "
     "lightweight and sufficient way to manage and share global state without the additional boilerplate of "
     "an external library. The application defines several contexts — for authentication, theme, toast "
     "notifications and onboarding — each wrapping the component tree and making its data available "
     "wherever it is needed.")
body("Figure 9 shows a context provider that stores the authenticated user and exposes login and logout "
     "functions to the entire component tree.")
figure("c6_context_provider.png", "A React Context Provider", width=5.7)

h3("3.3.1 Reducer Pattern for State Logic")
body("The reducer pattern is a common technique for keeping state transitions predictable: a reducer is a "
     "pure function that receives the current state and an action and returns the next state, centralising "
     "all state-updating logic in one place. For the relatively simple global state required by this "
     "project — the authenticated user, the selected theme and notifications — the Context API is used "
     "together with the useState hook, which is sufficient; the reducer pattern (via the useReducer hook) "
     "can be introduced for any context whose logic later grows more complex.")
h3("3.3.2 Context Provider as the Application Store")
body("A context provider component holds the shared state and exposes it, together with the functions that "
     "modify it, to all of its descendant components. The provider therefore plays the same role as a "
     "store: it is the single place where a particular slice of global state lives, and any component can "
     "read from or update it through the corresponding hook.")

body("The overall flow of state through the Context API is summarised in Figure 10.")
figure("c7_context_flow.png", "Context API State Management", width=6.2)

h2("3.4 Axios")
body("Axios is a promise-based HTTP client for JavaScript that runs in the browser and in Node.js. It is "
     "used on the client to send requests to the server’s REST API and to receive responses as JSON. "
     "In this project a single configured Axios instance attaches the authentication token to every request "
     "and centralises error handling, so that all of the application’s service modules communicate with "
     "the back end in a consistent way.")

h2("3.5 Application Programming Interface (API)")
body("An Application Programming Interface (API) is a defined set of rules that allows two software "
     "components to communicate. This project exposes a RESTful API: the server publishes a collection of "
     "endpoints, each identified by a URL and an HTTP method (GET, POST, PUT, DELETE), and the client calls "
     "these endpoints to read and modify data. All endpoints are versioned under the path /api/v1 and "
     "exchange data in JSON format. Representational State Transfer (REST) is an architectural style that "
     "treats server data as resources and uses standard HTTP verbs to operate on them, which makes the API "
     "predictable and easy to consume.")

h2("3.6 Cloud Database — MongoDB Atlas")
body("The application stores its data in MongoDB, which can be hosted locally during development or in the "
     "cloud using MongoDB Atlas. Atlas is the official cloud database service for MongoDB; it provides a "
     "managed, always-available database with automatic backups and security controls, accessed through a "
     "single connection string. Hosting the database in the cloud allows the application to be deployed and "
     "accessed from anywhere, and removes the burden of maintaining database server infrastructure.")

h2("3.7 Node JS")
body("Node.js is an open-source, cross-platform JavaScript runtime built on Google Chrome’s V8 engine. "
     "It allows JavaScript to be executed outside the browser and is used to run the server side of the "
     "application. Node.js uses an event-driven, non-blocking input/output model, which makes it efficient "
     "and well suited to data-intensive applications that handle many simultaneous requests.")
body("Figure 11 shows a minimal web server written using the built-in http module of Node.js.")
figure("c8_node_server.png", "A Simple Node.js Web Server", width=5.7)

h3("3.7.1 Features of Node JS")
bullet("Asynchronous and event-driven|| — operations do not block the program while waiting for "
       "input/output, allowing many requests to be handled concurrently.")
bullet("Single-threaded but scalable|| — a single thread with an event loop serves a large number of "
       "connections efficiently.")
bullet("Rich package ecosystem|| — the Node Package Manager (npm) provides access to a vast library "
       "of reusable open-source modules.")
bullet("Cross-platform|| — the same code runs on Windows, macOS and Linux.")

h2("3.8 Express JS")
body("Express.js is a fast, minimal and flexible web-application framework for Node.js. It simplifies the "
     "creation of a web server and a REST API by providing a clean way to define routes, apply middleware "
     "and handle requests and responses. In this project Express (version 5) is used to build the entire "
     "back-end API, with middleware for authentication, request validation, logging, rate limiting and "
     "centralised error handling.")

body("The same web server is considerably simpler to build with Express, as shown in Figure 12.")
figure("c9_express_server.png", "An Express Server", width=5.7)

h2("3.9 MongoDB and Mongoose")
h3("3.9.1 MongoDB")
body("MongoDB is a NoSQL, document-oriented database. Instead of storing data in tables of rows and "
     "columns as a relational database does, it stores data as flexible, JSON-like documents grouped into "
     "collections. This flexibility makes it easy to evolve the data model as the application grows, and "
     "its document structure maps naturally onto the objects used in JavaScript.")
body("In deployment the database is hosted on a managed MongoDB Atlas cluster, illustrated in Figure 13.")
figure("c10_mongo_cluster.png", "MongoDB Atlas Cluster", width=6.0)

h3("3.9.2 Key Components of MongoDB Architecture")
bullet("Document|| — the basic unit of data, stored in a binary JSON format called BSON.")
bullet("Collection|| — a group of related documents, analogous to a table in a relational database.")
bullet("Database|| — a container that holds one or more collections.")
bullet("Index|| — a structure that improves the speed of queries on selected fields.")
bullet("_id field|| — a unique identifier automatically assigned to every document, serving as its "
       "primary key.")
h3("3.9.3 Mongoose")
body("Mongoose is an Object Data Modeling (ODM) library for MongoDB and Node.js. It allows the developer "
     "to define schemas that describe the structure, data types and validation rules of documents, and "
     "then work with the data through model objects. In this project Mongoose defines all of the data "
     "models — User, Product, Sale, Customer, KhataEntry, Supplier, StockAdjustment and others — "
     "and enforces that monetary values are stored using the high-precision Decimal128 type.")

h2("3.10 Thunder Client / Postman")
body("Postman and Thunder Client are tools used to test REST APIs. They allow a developer to send HTTP "
     "requests to the server’s endpoints and inspect the responses without building a front end first. "
     "Thunder Client is a lightweight extension that runs inside Visual Studio Code. During development, "
     "these tools were used to test each endpoint of the API — verifying status codes, response "
     "structure and the correctness of the GST calculations — before connecting it to the React client.")

h2("3.11 Project Implementation")
body("The application is organised into two main parts: a client folder containing the React front end and "
     "a server folder containing the Node.js and Express back end. The server is further structured into "
     "routes, controllers, services, models, middlewares, validators and scheduled jobs, which keeps the "
     "code modular and maintainable. The client is structured into pages, reusable components, services "
     "that call the API, custom hooks and context providers. The two parts communicate exclusively through "
     "the REST API.")

body("The dependencies of the two parts of the application are declared in their package.json files. "
     "Figure 14 lists the principal client-side dependencies and Figure 15 the server-side dependencies, "
     "the latter including the AI, OCR, PDF and scheduling libraries that power the system.")
figure("c11_client_pkg.png", "Client-Side Dependencies (package.json)", width=5.7)
figure("c12_server_pkg.png", "Server-Side Dependencies (package.json)", width=5.7)

h2("3.12 Development Environment Setup")
body("The development environment was set up by installing Node.js (which includes the npm package "
     "manager), a code editor and access to a MongoDB database. The server and client dependencies are "
     "declared in their respective package.json files and installed with npm. During development the React "
     "client is served by the Vite development server with hot-module replacement, while the Express server "
     "runs as a separate Node.js process; the two communicate over HTTP.")

h2("3.13 Version Control System (VCS)")
body("A Version Control System records changes to the source code over time so that specific versions can "
     "be recalled later, and so that work can be managed in a disciplined way. This project uses Git, the "
     "most widely used distributed version control system. Git tracks the history of every file, allows "
     "changes to be grouped into commits with descriptive messages, and supports branching so that new "
     "features can be developed in isolation before being merged.")

h2("3.14 Visual Studio Code")
body("Visual Studio Code (VS Code) is a free, lightweight and extensible source-code editor developed by "
     "Microsoft. It offers features such as syntax highlighting, intelligent code completion, integrated "
     "debugging, Git integration and a large marketplace of extensions. VS Code was used as the primary "
     "development environment for writing, testing and debugging both the client and server code of this "
     "project.")

h2("3.15 JSON Web Tokens (JWT)")
body("A JSON Web Token (JWT) is a compact, self-contained and digitally signed token used to securely "
     "transmit information between parties. In this application JWTs are used for authentication. When a "
     "user logs in with valid credentials, the server verifies the password against its bcrypt hash and "
     "issues a signed token. The client stores this token and sends it with every subsequent request; the "
     "server’s authentication middleware verifies the signature before allowing access to any "
     "protected route. Because the token is self-contained, the server does not need to store session "
     "state, which keeps the API stateless and scalable. The authentication flow is shown in Figure 16.")
figure("d5_jwt.png", "JWT Authentication Flow", width=6.4)
page_break()

# ================================================================ CHAPTER 4
chapter_banner(4)
h1("4. Product Overview")
body("This chapter describes the system from the point of view of its users and requirements, and gives an "
     "overview of the client and server parts of the application.")

h2("4.1 Users and Stakeholders")
bullet("Business owner (primary user)|| — manages inventory, performs billing, tracks customer credit, "
       "and reviews analytics and AI insights to run the business.")
bullet("Shop staff|| — use the billing and inventory screens for day-to-day counter operations.")
bullet("Customers|| — benefit indirectly through faster billing, digital invoices, UPI payment and an "
       "accurate record of their credit.")
bullet("Suppliers|| — represented in the system through supplier records and purchase transactions.")
bullet("Accountant|| — consumes the Tally-compatible export of transaction data for bookkeeping and "
       "tax filing.")

h2("4.2 Functional and Non-Functional Requirements")
h3("4.2.1 Functional Requirements")
bullet("User registration, login and authentication with secure password storage.")
bullet("Creation and management of products, including category, price, GST rate, unit of measure and "
       "reorder level.")
bullet("Point-of-sale billing with automatic GST computation and PDF invoice generation.")
bullet("Recording of payments by cash, UPI/online or credit, with credit posted to the customer ledger.")
bullet("Management of customers and a digital khata with running balances and repayments.")
bullet("Management of suppliers and recording of stock adjustments (shrinkage).")
bullet("Demand forecasting and reorder suggestions derived from sales history.")
bullet("An AI help assistant (Google Gemini) and automatic reading of supplier bills (Gemini Vision "
       "with a Tesseract.js OCR fallback).")
bullet("Analytics dashboard, scheduled report generation and smart stock alerts.")
h3("4.2.2 Non-Functional Requirements")
bullet("Security|| — authenticated and validated APIs, hashed passwords and protection against abuse.")
bullet("Reliability|| — accurate monetary calculations using high-precision decimal arithmetic.")
bullet("Usability|| — a simple, multi-language interface suitable for non-technical users.")
bullet("Performance|| — a responsive single-page interface that updates efficiently.")
bullet("Maintainability|| — a modular, layered codebase that is easy to extend and test.")
bullet("Portability|| — a web application accessible from any modern browser.")

h2("4.3 Frontend Application (Client Side)")
body("The client side is a single-page application built with React. It presents the user interface "
     "through a set of pages — Dashboard, Inventory, Quick Sale, Customers and Khata, Suppliers, AI "
     "Insights, Analytics, Scanner, Reports and Settings — organised within a common dashboard layout "
     "with a sidebar and top navigation. The client communicates with the server entirely through the REST "
     "API using Axios, and manages shared state such as the authenticated user and theme through React "
     "context providers.")

h2("4.4 Backend Server (Server Side)")
body("The server side is a REST API built with Node.js and Express. It receives requests from the client, "
     "authenticates and validates them, executes the required business logic and returns JSON responses. "
     "The server also hosts the integrations that give the system its intelligence and automation: an "
     "AI help-assistant and insights module, a bill-reading service (Gemini Vision with a Tesseract.js "
     "OCR fallback), the PDF-invoice service and the scheduled cron jobs. "
     "It persists all data to MongoDB through Mongoose. The principal collections of the database and the "
     "relationships between them are shown in Figure 17.")
figure("d3_database.png", "Database Schema and Relationships", width=6.4)

h2("4.5 Client-Side Modules and Pages")
bullet("Authentication and onboarding|| — login, registration and a guided multi-step onboarding "
       "wizard for new businesses.")
bullet("Inventory pages|| — product listing, creation and editing, low-stock indicators and the "
       "shrinkage (stock-adjustment) screen.")
bullet("Quick Sale page|| — product search, weight or amount entry, the bill cart and the payment-mode "
       "picker.")
bullet("Customers and Khata pages|| — customer directory and the digital credit ledger.")
bullet("AI Insights and Analytics pages|| — demand forecasts, reorder report and charts rendered with "
       "Recharts.")
bullet("Scanner page|| — upload of a supplier bill for OCR-based product import.")

h2("4.6 Server-Side API and Services")
bullet("Route groups|| — modular endpoints for authentication, products, sales, customers, khata, "
       "suppliers, inventory adjustments, AI, OCR, analytics, reports and settings under /api/v1.")
bullet("Controllers and services|| — controllers handle the HTTP layer while services contain the "
       "business logic and database access.")
bullet("Middleware|| — JWT authentication, Zod-based validation, rate limiting and central error "
       "handling.")
bullet("Integration services|| — an AI help-assistant and insights module (rule-based analytics plus a "
       "Gemini-powered chatbot), a bill-reading service (Gemini Vision with a Tesseract.js OCR fallback), "
       "the PDFKit invoice service and the Tally export service.")
bullet("Scheduled jobs|| — a report generator and a smart-alerts job that run automatically using "
       "node-cron.")
page_break()

# ================================================================ CHAPTER 5
chapter_banner(5)
h1("5. Conclusion and Future Work")

h2("5.1 Critical Evaluation")
body("The AI Smart Inventory Management System successfully digitises the complete daily workflow of a "
     "small retail business — inventory, billing, customer credit, supplier management and analysis "
     "— within a single coherent application. Measured against the objectives set out at the start of "
     "the project, the system meets each of them: inventory is tracked in real time, billing applies GST "
     "automatically and produces a professional invoice, customer credit is maintained digitally, and "
     "artificial intelligence is used in a practical rather than decorative way.")
body("A particular strength of the implementation is its engineering discipline: monetary values use "
     "high-precision decimal arithmetic to eliminate rounding errors, every API request is validated "
     "against a schema, access is protected by token-based authentication, and invoice numbers are "
     "generated atomically to remain sequential and gap-free. The layered architecture keeps the codebase "
     "modular and testable. The main limitations of the current version are its dependence on an internet "
     "connection, the absence of a dedicated mobile application, and support for only a single business "
     "account per login — each of which is addressed in the future-work plan below.")

h2("5.2 Future Work")
reset_num()
numbered("Mobile application:|| a React Native or progressive web application so that billing can be "
         "performed on any phone at the counter.")
numbered("Messaging integration:|| automatic delivery of invoices and credit-payment reminders through "
         "WhatsApp or SMS.")
numbered("Barcode scanning:|| camera-based barcode support for rapid billing of packaged goods.")
numbered("Multi-store and multi-user support:|| workspaces with role-based permissions for staff members "
         "across multiple branches.")
numbered("Offline-first operation:|| local storage with background synchronisation so that billing "
         "continues during internet outages.")
numbered("Advanced intelligence:|| price optimisation, festival-season demand models and anomaly "
         "detection for identifying possible theft.")

h2("5.3 Final Thoughts")
body("The project demonstrates the full breadth of modern full-stack engineering: requirement analysis, "
     "layered system architecture, document database design, secure and validated REST APIs, a "
     "contemporary React user interface, the integration of artificial-intelligence and OCR services, "
     "scheduled automation, and testing. Beyond the academic exercise, the system addresses a genuine "
     "problem faced by millions of small retailers, and its modular design provides a solid foundation on "
     "which the enhancements described above can be built. The development of this project has been a "
     "valuable opportunity to apply theoretical knowledge to a complete, real-world software product.")
page_break()

# ================================================================ REFERENCES
front_title("REFERENCES")
refs = [
    "React – A JavaScript Library for Building User Interfaces. https://react.dev",
    "Express.js – Fast, Unopinionated, Minimalist Web Framework for Node.js. https://expressjs.com",
    "MongoDB Documentation. https://www.mongodb.com/docs",
    "Mongoose ODM Documentation. https://mongoosejs.com",
    "Node.js Documentation. https://nodejs.org/en/docs",
    "Google Gemini API (Google AI for Developers). https://ai.google.dev",
    "Tesseract.js – Pure JavaScript OCR. https://tesseract.projectnaptha.com",
    "PDFKit – A JavaScript PDF Generation Library. https://pdfkit.org",
    "Tailwind CSS Documentation. https://tailwindcss.com",
    "JSON Web Tokens. https://jwt.io",
    "Vite – Next Generation Frontend Tooling. https://vitejs.dev",
    "Gujarat Technological University (GTU). https://www.gtu.ac.in",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.first_line_indent = Inches(-0.4)
    p.add_run("[%d]  " % i).bold = True
    p.add_run(r)

doc.save(OUT)
print("SAVED:", OUT)
