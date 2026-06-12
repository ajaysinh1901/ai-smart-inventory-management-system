"""Build AI_Smart_Inventory_System.docx from CEO-drafted content."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CONTENT = r"""# AI Smart Inventory System

**India's first AI-native billing, inventory, and khata platform — built for the kirana, the pharmacy, and the distributor who run their day on WhatsApp, UPI, and a paper bahi.**

**Prepared by:** Zimbs Tech
**Date:** 29 April 2026
**Confidentiality:** This document contains commercially sensitive information. It is shared in confidence with the named recipient and must not be circulated, copied, or referenced externally without prior written consent from Zimbs Tech.

---

## Executive Summary

AI Smart Inventory System is a billing, inventory, and credit-ledger platform purpose-built for Indian small and medium businesses — kirana stores, pharmacies, distributors, wholesalers, and online sellers doing roughly Rs. 40 lakh to Rs. 10 crore in annual turnover. We replace the three things every Indian shopkeeper juggles separately today — a paper bill book, a Tally-style inventory register, and a khata notebook — with one modern, Hindi-first application that does GST invoicing, runs the customer credit ledger, ingests supplier invoices via OCR, and answers business questions through a Gemini-powered AI Copilot.

The product is no longer a slide. **Phase 1 MVP is locked as of April 2026** and runs end-to-end: 47 REST endpoints across 13 domains, 10 MongoDB collections, gap-free fiscal-year invoice numbering, atomic stock deduction, an append-only khata ledger, OCR goods-receipt flow, daily smart-alerts cron, and a streaming AI Copilot with live tool access to the customer's own data.

The market is moving in our direction. The GST e-invoicing turnover threshold has dropped to Rs. 5 crore and is widely expected to fall further within twelve months — every Vyapar customer above that line becomes a forced-migration candidate the moment their CA sends the email. UPI now carries roughly seventy percent of retail payment volume. Gemini Flash made truly useful AI affordable to ship at SMB prices for the first time. None of the incumbents — Vyapar, Tally, Khatabook, Zoho Books, Marg — combine modern AI, OCR, and a Hindi-first interface in one app.

We are raising a seed round to fund eighteen months of focused growth across Pune, Indore, Ahmedabad, and Bengaluru. We are also opening a CA channel partnership and a thirty-day Growth-tier pilot for selected SMBs. Details in The Ask.

---

## The Problem

Indian SMBs are forced into formal accounting by GST but still run their day on a paper bill book, a WhatsApp group, and a Tally file that the accountant opens once a quarter. That gap costs them money every single day.

**Sunita, 45, runs a single-doctor pharmacy in Indore.** Her turnover is Rs. 1.4 crore a year across 4,500 SKUs. Sixty percent of her business is on credit — clinics nearby, monthly families, the occasional walk-in regular. She loses an estimated **Rs. 40,000 every month** to forgotten dues that she meant to chase but never did. She mistypes medicine names at the counter because she has no barcode scanner. She has no way to know which strip is about to expire until it already has.

**Mehul, 38, runs a kirana plus FMCG distribution business in Pune.** Turnover Rs. 6.5 crore, 4 staff, 2,800 SKUs, roughly 120 invoices a day. He is already above the e-invoicing threshold. On the 27th of every month his CA emails him asking *"GSTR-1 ready ya nahi?"* and he spends the next three days exporting from Vyapar, copy-pasting into Excel, and uploading to the GST portal by hand. When he raises an invoice on his phone, his stock register on his desktop does not update.

**Rohan, 29, sells D2C apparel in Bengaluru** across Shopify, Meesho, and a walk-in shop. Rs. 85 lakh a year. His inventory drifts across channels because nothing is the single source of truth, and his GSTR-1 reconciliation takes four hours every week.

These are not edge cases. They are the median customer in a market of roughly **8 to 10 million SMBs** doing Rs. 40 lakh to Rs. 10 crore in turnover. The pain is concrete, recurring, and measurable in rupees.

The incumbents do not solve it well. Vyapar owns the offline shopkeeper but has no real AI and no OCR. Tally owns the accountant but is unusable on a phone and visually frozen in 2008. Khatabook owns the credit-book mental model but has no GST invoice and quietly stepped back from lending. Zoho Books is great for the digitally native English-speaking seller but loses the Tier-2 owner the moment the demo opens. Nobody has shipped a Hindi-first, AI-native, all-in-one app for this segment.

---

## The Solution: AI Smart Inventory System

**One app. GST invoices, customer khata, supplier OCR, smart alerts, and an AI Copilot that actually understands your shop. In Hindi or English. On a Rs. 12,000 phone.**

### GST Invoicing that an Indian buyer expects

Create a sale, pick the customer, and the system auto-detects intra-state versus inter-state and splits CGST/SGST or IGST correctly. HSN codes are frozen on the line item so historical invoices remain accurate even if the product is later edited. Invoice numbers are gap-free per fiscal year (INV-2026-00001), allocated atomically — no duplicate, no skipped number, ever. The PDF is A4-clean, generated server-side, with the amount in words using Indian numbering (Crore, Lakh, Thousand). Every printed invoice carries a **UPI Scan-to-Pay QR** in the footer so a customer under thirty-five can pay before they leave the counter. WhatsApp share is one tap.

### Khata that remembers what you forgot

Every sale automatically posts a debit to the customer's ledger. Payments — cash, UPI, cheque, or bank — post a credit with a receipt number. The ledger is **append-only**: corrections become reversal entries with a reference to the original, so the audit trail stays intact for your CA. A Top Debtors view shows you exactly who to call today. Customer statements export as JSON or PDF. This is the feature Sunita pays us for.

### OCR Goods Receipt — stop typing supplier invoices

Drag a photo of your supplier's invoice into the Scanner. Tesseract reads it, our parser pulls invoice number, vendor, line items, and totals, and you get a review screen to correct anything OCR got wrong. One confirmation creates the GRN, increments stock atomically, and writes a stock-IN audit row linked to the original image. What used to be twenty minutes of typing becomes ninety seconds of review.

### AI Copilot you can actually ask in plain English (or Hindi)

Open the Copilot and type *"what should I reorder this week?"* or *"who owes me the most money?"*. The model has typed tool access to the customer's own database — get_low_stock, get_top_movers, get_dead_stock, get_gst_summary, get_supplier_list — and answers from live data, not a generic LLM hallucination. It also drives a 30-day demand forecast, dead-stock detection, and per-product reorder suggestions based on velocity and lead time.

### Smart Alerts that fire before you lose money

A daily cron at 9 AM IST scans every product for out-of-stock, low-stock (below your custom threshold), and dead-stock (no sale in thirty days, stock greater than zero). Alerts are de-duplicated, surfaced in the top-nav bell, and dismissable. Admins can trigger the scan on demand.

### A Dashboard built for skimmers

Four KPI tiles at the top — today's sales, GST collected this month, low-stock count, open alerts. A 30-day revenue line chart, a Top 5 products bar chart, recent transactions, and three Quick Action buttons. The four numbers that matter are visible in under two seconds.

### Hindi UI — not a translation, a first-class language

Every user-facing string flows through react-i18next with English and Hindi locales shipping today. The language switcher persists per user. The invoice PDF supports Devanagari. Marathi, Gujarati, Tamil, and Telugu are next.

---

## Target Customers

We sell to three personas, ranked by addressable volume.

### Persona A — Mehul, 38 — Kirana plus FMCG distributor, Pune

| Attribute | Value |
|---|---|
| Turnover | Rs. 6.5 crore per year |
| Staff | 4 |
| SKUs | 2,800 |
| Invoices/day | ~120 |
| Today's tool | Vyapar Premium at Rs. 3,599/year |
| Top pain | E-invoicing upload to GST portal; stock not auto-deducting when he bills on his phone |
| Why he pays us | E-invoicing built-in (Phase 2), AI reorder copilot, Tally export, GSTR-1 export |
| Annual ARPU | Rs. 4,800 to Rs. 6,000 |
| 3-year LTV | ~Rs. 15,000 |

### Persona B — Sunita, 45 — Single-doctor pharmacy, Indore

| Attribute | Value |
|---|---|
| Turnover | Rs. 1.4 crore per year |
| Staff | 1 helper |
| SKUs | 4,500 |
| Credit business | 60% |
| Top pain | Loses ~Rs. 40,000/month to forgotten dues; mistypes medicine names |
| Why she pays us | Khata ledger with WhatsApp reminders, OCR GRN, barcode scan and expiry tracking (Phase 1.5) |
| Annual ARPU | Rs. 2,400 to Rs. 3,600 |
| 3-year LTV | ~Rs. 9,000 |

This is our volume segment. There are an estimated 3 to 4 million Sunita-tier businesses in India.

### Persona C — Rohan, 29 — D2C apparel + Instagram seller, Bengaluru

| Attribute | Value |
|---|---|
| Turnover | Rs. 85 lakh per year |
| Channels | Shopify + Meesho + walk-in |
| Today's tool | Zoho Books at Rs. 6,000/year |
| Top pain | Inventory drift across channels; GSTR-1 reconciliation hell |
| Why he pays us | Unified inventory + AI insights, single source of truth |
| Annual ARPU | Rs. 6,000 to Rs. 9,000 |

Smaller TAM but the highest ARPU and the loudest on social media — strategic for word-of-mouth.

---

## Competitive Landscape

| Capability | AI Smart Inventory | Vyapar | Tally Prime | Khatabook | Zoho Books |
|---|---|---|---|---|---|
| Hindi + regional UI | Yes | Yes | Partial | Yes | English-first |
| GST invoice (CGST/SGST/IGST) | Yes | Yes | Yes | No | Yes |
| Khata ledger | Yes | Yes | DIY | Yes | Partial |
| OCR supplier invoice → GRN | Yes | No | No | No | No |
| AI Copilot with live tool-use | Yes | No | No | No | Basic |
| All-in-one (billing + khata + inventory) | Yes | Yes | Yes | Ledger only | Yes |
| Modern dark-mode UX | Yes | No | No | Partial | Partial |
| E-invoicing (IRN + signed QR) | Phase 2 | Add-on | Add-on | No | Yes |
| Annual price band (paid tier) | Rs. 4,999 | Rs. 3,599 | Rs. 18,000+ | Free | Rs. 6,000+ |

**Where we win clearly today:** OCR-driven GRN, AI Copilot with live tool access, modern UX, single-app coverage of billing + khata + inventory.

**Where we do not yet win, honestly:** E-invoicing (Phase 2, Q3 2026), Tally export at parity (planned), regional language coverage beyond Hindi (Phase 2). We are also younger than every name on this list and have to earn trust customer by customer.

---

## Pricing & Business Model

Indian SMBs hate per-seat pricing and love annual prepay with a real (not fake) discount. Our tiering is built around that reality.

| Tier | Price | Target customer | Key inclusions |
|---|---|---|---|
| Free | Rs. 0 forever | Single-shop kirana, hobbyists | Up to 50 invoices/month, 1 user, 200 SKUs, GST invoice, basic khata |
| Growth | Rs. 599/month or Rs. 4,999/year | Sunita-tier pharmacy/kirana | Unlimited invoices and SKUs, 3 users, UPI Scan-to-Pay, full khata + WhatsApp reminders, OCR GRN, AI Copilot |
| Pro | Rs. 9,600/year (indicative) | Mehul-tier distributor | Everything in Growth + e-invoicing (IRN/QR), GSTR-1/3B + Tally export, barcode + thermal print, 5 users, Hindi + regional UI, priority support |

### Why these numbers

Growth at Rs. 4,999/year is **39% above Vyapar Silver (Rs. 1,899/year)** and roughly at parity with Vyapar Premium (Rs. 3,599/year), justified by the Khata ledger + AI Copilot + OCR that none of them ship. Pro is priced **roughly 2.7x Vyapar Premium**, justified entirely by automating the manual GST-portal upload that Mehul currently spends three days a month on.

### Annual prepay discount

Twenty percent off annual prepay versus monthly equivalent. Indian SMBs sniff out fake discounts in a single demo, so the discount is real and stable.

### CA partner programme

Refer a paid customer and the CA earns **Rs. 400/customer/year recurring** for as long as the customer stays. CAs become our distribution channel; their referrals retain at over ninety percent because the CA has already done the trust work.

### Indicative unit economics

| Metric | Target |
|---|---|
| Blended ARPU | Rs. 3,000/year |
| Gross margin | ~80% (after Gemini API + WhatsApp + hosting) |
| CAC via CA channel | < Rs. 800 |
| CAC via direct/digital | Rs. 1,200 to Rs. 1,800 |
| Payback period | 4 to 7 months |
| Logo churn target | < 4% per month at 12-month mark |

---

## Product Status: What Is Shipped Today

This is a working product, not a deck. Everything below is in the codebase as of 29 April 2026.

### Phase 1 MVP — locked

- GST invoice with CGST/SGST/IGST auto-split based on seller-buyer state
- Gap-free per-fiscal-year invoice numbering, allocated atomically
- Per-line frozen HSN code (history-safe)
- Server-side A4 PDF with amount in words (Indian numbering)
- UPI Scan-to-Pay QR embedded on every invoice
- WhatsApp share with prefilled invoice text
- Customer master with GSTIN (15-char), phone (+91), and 36 Indian states/UTs
- Append-only khata ledger with reversal-entry audit pattern
- Top Debtors view and PDF customer statements
- Inventory with SKU, HSN, GST rate, cost/selling price, reorder threshold, supplier link, soft delete
- OCR supplier-invoice ingest via Tesseract.js with mandatory human review
- Stock movement audit log on every IN/OUT/ADJUST
- Smart Alerts daily cron (out-of-stock, low-stock, dead-stock)
- AI Insights Copilot with five live tool functions, streaming responses, suggested-question chips
- 30-day demand forecast, per-product reorder suggestion, dead-stock list, sales-velocity trends
- Dashboard with 4 KPI tiles, 30-day revenue chart, top-5 products, recent transactions
- Analytics page (revenue trend, category breakdown, inventory health, profit analysis)
- Settings (workspace + GSTIN + UPI + AI config + notifications + user management)
- 60-second onboarding wizard for new admins
- Role-based access control (admin / manager / staff)
- English + Hindi UI with persistent language switcher
- Dark mode
- Tally Prime XML export endpoint

### Phase 1.5 — next 30 days

- Camera-based barcode scan via @zxing/browser (Product schema is barcode-ready)
- 58 mm thermal-printer invoice template (browser print, ESC/POS optimised)
- Marathi, Gujarati, Tamil, Telugu UI rolling out behind native-speaker review

### Phase 2 — Q3 2026

- E-invoicing (IRN generation, signed QR, AckNo/AckDate) via a sandbox-certified GSP — ClearTax, Masters India, or IRIS
- WhatsApp Business API for automated khata payment reminders (replaces the current wa.me deep-link)
- GSTR-1 and GSTR-3B JSON export so the customer's CA stays inside our app
- Purchase Order workflow with PO-to-GRN matching
- Multi-location and multi-warehouse with consolidated reporting
- Per-tenant token accounting on the AI Copilot (prerequisite for multi-tenant SaaS)

---

## Technology Snapshot

We chose a stack we can ship on, hire for, and operate cheaply at SMB scale.

### Stack

- **Client:** React 19 + Vite 8 + Tailwind 3, lazy-loaded routes, code-split bundle under 300 kB initial chunk
- **Server:** Node 20 + Express 5, versioned REST at /api/v1, 47 endpoints across 13 domains
- **Database:** MongoDB with Mongoose 9, 10 collections, atomic counter for sequence allocation
- **AI:** Google Gemini 2.5 Flash with typed tool-use; 20 requests per minute per IP rate limit to keep API costs predictable
- **OCR:** Tesseract.js with a heuristic parser tuned for Indian invoice layouts
- **Localisation:** react-i18next with English and Hindi locales today
- **PDF:** server-side pdfkit, A4 format, dynamic tax columns

### Security posture

- Bcrypt password hashing (10 rounds), JWT (HS256) with 7-day expiry
- Zod schema validation on every POST/PUT route — clean 400s with field-level errors, never a 500 stack
- Format guards: GSTIN (15-char regex), Indian phone (+91 with 10 digits), HSN (4 to 8 digits), state enum (36 Indian states/UTs)
- Rate limits: auth 15/min, AI chat 20/min, global 200/min per IP
- **Atomic invariants** that matter for an accounting product:
  - Invoice numbers allocated via findOneAndUpdate with $inc on a Counter — race-free, gap-free
  - Stock deduction conditional on stock >= qty — atomic oversell prevention
  - Khata reversal never mutates the original entry — append-only audit trail

### Why this matters to a technical reader

We did not pick the trendy stack. We picked the boring stack and got the invariants right. An accounting application that allows duplicate invoice numbers, negative stock, or mutable ledger entries is unsellable to a CA — and CAs decide what their clients buy. Those guards are in production today.

---

## Roadmap

### 0 to 3 months (May–July 2026)

- Barcode scan-to-sell across the catalog
- Thermal-printer invoice template (58 mm)
- Marathi, Gujarati, Tamil, Telugu UI live
- Pricing page and paywall enforcement
- Multi-tenant data isolation with workspace scoping
- Razorpay subscription billing on the existing tier model
- CA partner programme launches with 10 hand-picked CAs in Pune, Indore, Ahmedabad
- **Target: 200 paying customers, ~Rs. 6 lakh MRR**

### 3 to 6 months (August–October 2026)

- E-invoicing (IRN + signed QR) certified through a GSP
- WhatsApp Business API for automated khata reminders
- GSTR-1 and GSTR-3B JSON export
- Purchase Order workflow with PO-to-GRN matching
- Multi-branch and multi-warehouse with consolidated reporting
- Customer mobile app (React Native, read-mostly) for owners on the move
- **Target: 500 paying customers, ~Rs. 15 lakh MRR**

### 6 to 12 months (November 2026–April 2027)

- Marketplace integrations: Shopify, Amazon, Meesho (unlocks Persona C at scale)
- Reverse-charge mechanism, composite supply, exempt items — full GST edge-case coverage
- ABC analysis, stock aging, and slow-mover reports
- Public REST API and webhooks for accountants and ERP integrators
- Distributor-tier features: route-based ordering, schemes and target tracking
- **Target: 2,000 paying customers, ~Rs. 60 lakh MRR**

---

## Why Now

**Regulatory tailwind.** The GST e-invoicing turnover threshold has dropped from Rs. 500 crore to Rs. 5 crore in five years. Industry expectation is that it falls to Rs. 1–2 crore within the next twelve months. Every drop is a forced-migration moment for a band of SMBs whose CA tells them their current tool will not work next quarter. We need to be on the shortlist when that email goes out.

**Payment rail tailwind.** UPI now carries roughly seventy percent of retail payment volume in India. A printed invoice without a Scan-to-Pay QR feels archaic to any buyer under thirty-five. We ship the QR by default; most incumbents either do not, or charge for it as a Premium add-on.

**AI cost collapse.** Gemini 2.5 Flash brought the cost of a useful, tool-using AI assistant down to a level where we can include it in a Rs. 4,999/year SaaS plan without losing money. Two years ago this product was financially impossible at SMB pricing. Today it is a feature the incumbents cannot easily clone because their architectures predate function-calling.

**Hindi-first SaaS gap.** Vyapar's thirty million downloads were built largely on the strength of being Hindi-first. Tally, Zoho, and most modern SaaS still treat Hindi as a translation afterthought. We treated it as a first-class language from day one, including the invoice PDF.

**Distribution rail.** WhatsApp Business API pricing has stabilised at roughly Rs. 0.30 to Rs. 0.80 per utility message. Automated khata reminders are now economically viable to ship at our price point.

These four shifts only have to be true together, and they are, today.

---

## Team

**Founders**
- *[Founder name placeholder]* — CEO, Product
- *[Founder name placeholder]* — CTO, Engineering

**Advisors**
- *[Advisor name placeholder]* — GST and compliance
- *[Advisor name placeholder]* — Indian SMB distribution
- *[Advisor name placeholder]* — SaaS go-to-market

**Hiring plan (post-seed, 12-month horizon)**
- Two full-stack engineers
- One designer (mobile-first, Indic UX)
- One CA partner-success lead
- Two regional sales leads (Pune, Ahmedabad)

*[Founder bios and advisor names to be filled before final distribution.]*

---

## Traction & Pilot Status

- **MVP locked:** April 2026 — 47 endpoints, 10 collections, end-to-end demo trio working
- **Pilot pipeline:** Conversations underway with 3 Pune-area distributors and 2 Indore pharmacies for thirty-day Growth-tier trials in May 2026
- **CA channel:** Initial outreach to 12 CAs across Pune, Indore, and Ahmedabad; first partner programme cohort target of 10 by July 2026
- **Internal benchmarks:**
  - GST invoice creation in under 30 seconds end-to-end
  - OCR supplier-invoice review-and-confirm in under 90 seconds
  - AI Copilot first-token latency under 1.5 seconds on a Rs. 12,000 Android phone tethered to 4G

*[Live customer counts and revenue figures to be added once pilots convert in May 2026.]*

---

## The Ask

We are inviting three different readers to take three different actions. Pick the one that fits.

### To prospective investors

We are raising an **Rs. 3.5 crore seed round at an Rs. 18 crore pre-money valuation** for an 18-month runway. Use of funds: two engineering hires, one designer, two regional sales leads, GSP onboarding fees and certification, WhatsApp Business API onboarding, and roughly Rs. 60 lakh of working capital reserved for performance marketing in Pune, Indore, Ahmedabad, and Bengaluru.

Target on milestone money: **2,000 paying customers, Rs. 60 lakh MRR, blended ARPU Rs. 3,000/year, logo churn under 4%/month** by month 12 of the round.

We would value investors who bring SMB distribution, CA-network access, or GSP/IRP relationships in addition to capital.

### To channel partners (CAs and accounting firms)

Pilot **AI Smart Inventory System with five SMBs in your client network** at no cost to them on the Free tier, with full Growth-tier features unlocked for sixty days. We provide onboarding, the Hindi UI, and direct support. You earn **Rs. 400/customer/year recurring** for every client who upgrades to Growth or Pro and stays.

Your client gets a working GST + khata + AI tool. You get a recurring referral stream and one fewer client calling you on the 27th. We get a distribution channel that retains at ninety percent.

### To pilot SMB customers (Mehul-tier and Sunita-tier)

Try **Growth tier free for 30 days** in exchange for one thirty-minute feedback call and consent for a written case study (anonymised on request). No credit card. No auto-billing. Bring your real invoices, your real khata, and your real supplier bills. If we have not saved you measurable time or money in those thirty days, walk away — and if you stay, your first year is at a 25 percent founding-customer discount on the published price.

---

## Contact

**[Founder name placeholder]**
Founder & CEO, Zimbs Tech
Email: developers@zimbstech.com
Phone: *[phone placeholder]*
Website: *[website placeholder]*

For investor diligence requests, partnership conversations, or pilot enrolment, please reply to the email above with the subject line *"AI Smart Inventory — [Investor / Partner / Pilot]"*.

---

*This document reflects the AI Smart Inventory System product and codebase as of 29 April 2026. Pricing, roadmap dates, and target metrics are forward-looking and subject to change. Numbers cited for incumbent products are publicly available figures at the time of writing.*
"""


# ---------- helpers ----------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_inline_runs(paragraph, text):
    """Render markdown inline (**bold**, *italic*) as runs."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*\n]+\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        else:
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '8B1E1E')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def parse_table_block(lines, idx):
    """Consume a contiguous markdown table starting at lines[idx]; return (rows, next_idx)."""
    rows = []
    while idx < len(lines) and '|' in lines[idx] and lines[idx].strip().startswith('|'):
        rows.append([c.strip() for c in lines[idx].strip().strip('|').split('|')])
        idx += 1
    # Drop the separator row (---)
    cleaned = [r for r in rows if not all(set(c) <= set('-: ') for c in r)]
    return cleaned, idx


# ---------- build doc ----------

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Heading colors
PRIMARY = RGBColor(0x8B, 0x1E, 0x1E)  # Bahi Red
INK = RGBColor(0x14, 0x11, 0x0D)
BRASS = RGBColor(0xC8, 0x97, 0x3F)

for level, size in [(1, 26), (2, 18), (3, 13)]:
    s = doc.styles[f'Heading {level}']
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = PRIMARY if level <= 2 else INK

lines = CONTENT.split('\n')
i = 0
title_done = False

while i < len(lines):
    raw = lines[i]
    line = raw.rstrip()

    # Title page
    if not title_done and line.startswith('# '):
        # Big centered title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(4):
            p.add_run('\n')
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(line[2:])
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        title_done = True
        i += 1
        # Pull tagline (next non-empty line)
        while i < len(lines) and not lines[i].strip():
            i += 1
        if i < len(lines) and lines[i].startswith('**'):
            tagline = doc.add_paragraph()
            tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tag_text = lines[i].strip().strip('*')
            tr = tagline.add_run(tag_text)
            tr.font.size = Pt(14)
            tr.font.italic = True
            tr.font.color.rgb = INK
            i += 1
        # Skip blank
        while i < len(lines) and not lines[i].strip():
            i += 1
        # Prepared-by / Date / Confidentiality block
        for _ in range(8):
            doc.add_paragraph()
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('---'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, lines[i].strip())
            for run in p.runs:
                run.font.size = Pt(11)
            i += 1
        # Page break after cover
        doc.add_page_break()
        continue

    # Horizontal rule
    if line.strip() == '---':
        add_horizontal_rule(doc)
        i += 1
        continue

    # Headings
    if line.startswith('## '):
        doc.add_heading(line[3:], level=1)
        i += 1
        continue
    if line.startswith('### '):
        doc.add_heading(line[4:], level=2)
        i += 1
        continue

    # Tables (start with |)
    if line.strip().startswith('|') and '|' in line.strip()[1:]:
        rows, new_i = parse_table_block(lines, i)
        if rows:
            tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = True
            tbl.style = 'Light Grid Accent 1'
            for r_idx, row in enumerate(rows):
                cells = tbl.rows[r_idx].cells
                for c_idx, val in enumerate(row):
                    if c_idx >= len(cells):
                        break
                    cell = cells[c_idx]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    add_inline_runs(p, val)
                    if r_idx == 0:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        set_cell_shading(cell, '8B1E1E')
                    else:
                        for run in p.runs:
                            run.font.size = Pt(10)
            doc.add_paragraph()
        i = new_i
        continue

    # Bullet list
    if line.lstrip().startswith('- '):
        indent = len(line) - len(line.lstrip())
        text = line.lstrip()[2:]
        p = doc.add_paragraph(style='List Bullet')
        if indent >= 2:
            p.paragraph_format.left_indent = Cm(0.8)
        add_inline_runs(p, text)
        i += 1
        continue

    # Blank
    if not line.strip():
        i += 1
        continue

    # Regular paragraph
    p = doc.add_paragraph()
    add_inline_runs(p, line)
    i += 1


out = r'c:/Users/Admin/Desktop/Clg Mern/AI_Smart_Inventory_System.docx'
doc.save(out)
print(f'Saved: {out}')
