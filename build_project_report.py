"""Build AI_Smart_Inventory_System_Project_Report.docx — college submission format."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

STUDENT_NAME = "Ajaysinh Parmar"
PROJECT_TITLE = "AI Smart Inventory System"
SUBTITLE = "An AI-Powered Inventory Management, GST Billing & Customer Ledger Platform"
YEAR = "2025 - 2026"

# ===================== helpers =====================

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def add_centered(doc, text, size=12, bold=False, italic=False, color=None, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def add_para(doc, text, size=11, bold=False, italic=False, justify=True, indent=0):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x14, 0x11, 0x0D)
    return h


def add_chapter_heading(doc, num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"CHAPTER {num}")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x8B, 0x1E, 0x1E)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    r2 = p2.add_run(title.upper())
    r2.font.size = Pt(20)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x14, 0x11, 0x0D)
    add_horizontal_rule(doc)


def add_section_heading(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{num}  {title}")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x8B, 0x1E, 0x1E)


def add_subsection_heading(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{num}  {title}")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x14, 0x11, 0x0D)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.font.size = Pt(11)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '8B1E1E')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_table(doc, headers, rows, col_widths_cm=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Light Grid Accent 1'
    # Header row
    for c_idx, h in enumerate(headers):
        cell = tbl.rows[0].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '8B1E1E')
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    if col_widths_cm:
        for row in tbl.rows:
            for c_idx, w in enumerate(col_widths_cm):
                if c_idx < len(row.cells):
                    row.cells[c_idx].width = Cm(w)
    doc.add_paragraph()
    return tbl


def add_placeholder_box(doc, label):
    """Insert a single-cell bordered table acting as a screenshot placeholder."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(15)
    cell.text = ''
    for _ in range(4):
        cell.add_paragraph()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ {label} ]")
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    for _ in range(3):
        cell.add_paragraph()
    # Add border
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'dashed')
        b.set(qn('w:sz'), '8')
        b.set(qn('w:color'), '888888')
        tc_borders.append(b)
    tc_pr.append(tc_borders)
    doc.add_paragraph()


# ===================== build =====================

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)  # extra left margin for binding
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

for level, size in [(1, 18), (2, 14), (3, 12)]:
    s = doc.styles[f'Heading {level}']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(size)
    s.font.bold = True

PRIMARY = RGBColor(0x8B, 0x1E, 0x1E)
INK = RGBColor(0x14, 0x11, 0x0D)

# ============ TITLE PAGE ============
add_centered(doc, "A PROJECT REPORT", size=14, bold=True, space_before=24)
add_centered(doc, "ON", size=12, bold=False, space_before=6)
doc.add_paragraph()
add_centered(doc, f'"{PROJECT_TITLE}"', size=22, bold=True, color=PRIMARY)
add_centered(doc, SUBTITLE, size=12, italic=True)
doc.add_paragraph()
doc.add_paragraph()
add_centered(doc, "Submitted in partial fulfilment of the requirements", size=12)
add_centered(doc, "for the award of the degree of", size=12)
doc.add_paragraph()
add_centered(doc, "BACHELOR OF ENGINEERING", size=15, bold=True)
add_centered(doc, "in", size=11)
add_centered(doc, "COMPUTER ENGINEERING", size=14, bold=True)
doc.add_paragraph()
doc.add_paragraph()
add_centered(doc, "Submitted by", size=12)
add_centered(doc, STUDENT_NAME.upper(), size=18, bold=True, color=PRIMARY)
add_centered(doc, "[Enrollment Number]", size=11, italic=True)
doc.add_paragraph()
doc.add_paragraph()
add_centered(doc, "Under the guidance of", size=12)
add_centered(doc, "[Guide Name]", size=14, bold=True)
add_centered(doc, "[Designation, Department of Computer Engineering]", size=11, italic=True)
doc.add_paragraph()
doc.add_paragraph()
add_centered(doc, "[NAME OF DEPARTMENT]", size=13, bold=True)
add_centered(doc, "[NAME OF COLLEGE / INSTITUTE]", size=13, bold=True)
add_centered(doc, "[Affiliated University]", size=11, italic=True)
add_centered(doc, f"Academic Year: {YEAR}", size=12, bold=True, space_before=12)
add_page_break(doc)

# ============ CERTIFICATE ============
add_centered(doc, "CERTIFICATE", size=22, bold=True, color=PRIMARY, space_before=24)
add_horizontal_rule(doc)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 2.0
p.paragraph_format.first_line_indent = Cm(1.0)
p.add_run("This is to certify that the project report entitled ")
r = p.add_run(f'"{PROJECT_TITLE}"')
r.font.bold = True
p.add_run(" has been successfully carried out by ")
r = p.add_run(STUDENT_NAME)
r.font.bold = True
p.add_run(" (Enrollment Number: ____________________) under my supervision and guidance, "
          "in partial fulfilment of the requirements for the award of the degree of "
          "Bachelor of Engineering in Computer Engineering during the academic year " + YEAR + ".")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 2.0
p.paragraph_format.first_line_indent = Cm(1.0)
p.add_run("To the best of my knowledge, the work presented in this report is original and "
         "has not been submitted elsewhere for the award of any other degree or diploma.")

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Signature block (table)
sig = doc.add_table(rows=2, cols=3)
sig.autofit = True
sig_cells = sig.rows[0].cells
sig_cells[0].text = '\n_______________________'
sig_cells[1].text = '\n_______________________'
sig_cells[2].text = '\n_______________________'
sig_cells2 = sig.rows[1].cells
sig_cells2[0].text = 'Project Guide'
sig_cells2[1].text = 'Head of Department'
sig_cells2[2].text = 'Principal'
for row in sig.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()
add_centered(doc, "Place: ____________________", size=11)
add_centered(doc, "Date:  ____________________", size=11)
add_page_break(doc)

# ============ DECLARATION ============
add_centered(doc, "DECLARATION", size=22, bold=True, color=PRIMARY, space_before=24)
add_horizontal_rule(doc)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 2.0
p.paragraph_format.first_line_indent = Cm(1.0)
p.add_run("I, ")
r = p.add_run(STUDENT_NAME)
r.font.bold = True
p.add_run(", hereby declare that the project work entitled ")
r = p.add_run(f'"{PROJECT_TITLE}"')
r.font.bold = True
p.add_run(" submitted to [University Name] in partial fulfilment of the requirements "
          "for the award of the degree of Bachelor of Engineering in Computer Engineering "
          "is a record of original work carried out by me under the guidance and supervision "
          "of [Guide Name].")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 2.0
p.paragraph_format.first_line_indent = Cm(1.0)
p.add_run("I further declare that the matter embodied in this report has not been "
         "submitted by me to any other university or institute for the award of any "
         "degree, diploma, or similar title. All sources of information used in this "
         "report have been duly acknowledged in the references section.")

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("\n").font.size = Pt(11)
r = p.add_run(STUDENT_NAME)
r.font.bold = True
r.font.size = Pt(13)
p.add_run("\n[Enrollment Number]\n").font.size = Pt(11)
p.add_run("Department of Computer Engineering").font.size = Pt(11)

add_centered(doc, "Place: ____________________", size=11, space_before=18)
add_centered(doc, "Date:  ____________________", size=11)
add_page_break(doc)

# ============ ACKNOWLEDGEMENT ============
add_centered(doc, "ACKNOWLEDGEMENT", size=22, bold=True, color=PRIMARY, space_before=24)
add_horizontal_rule(doc)
doc.add_paragraph()

paragraphs = [
    "The completion of this project would not have been possible without the support, "
    "encouragement, and contributions of many individuals. I take this opportunity to "
    "express my sincere gratitude to all those who have helped me throughout this journey.",

    "First and foremost, I would like to express my deepest gratitude to my project guide, "
    "[Guide Name], for the constant guidance, valuable insights, and continuous encouragement "
    "throughout the duration of this project. Their expertise and patience were instrumental "
    "in shaping the direction and quality of this work.",

    "I am also grateful to [Head of Department], Head of the Department of Computer Engineering, "
    "for providing me with the opportunity, infrastructure, and academic environment to "
    "undertake this project. My sincere thanks to all the faculty members of the department "
    "for their support and timely feedback during project reviews.",

    "I would like to thank the Principal, [Principal Name], for fostering a research-oriented "
    "atmosphere in the institute and for the encouragement extended to project work.",

    "I am thankful to the open-source community, whose libraries and tools — including "
    "MongoDB, Express, React, Node.js, Tailwind CSS, Tesseract.js, and the Google Gemini "
    "API — formed the foundation of this implementation.",

    "Finally, I owe a debt of gratitude to my parents, family members, and friends for "
    "their unwavering moral support, patience, and encouragement throughout my engineering "
    "studies and during the demanding phases of this project.",
]

for para in paragraphs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(para)
    r.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run(STUDENT_NAME)
r.font.bold = True
r.font.size = Pt(13)
add_page_break(doc)

# ============ ABSTRACT ============
add_centered(doc, "ABSTRACT", size=22, bold=True, color=PRIMARY, space_before=24)
add_horizontal_rule(doc)
doc.add_paragraph()

abstract_paras = [
    "Small and medium-sized businesses in India face a daily struggle managing inventory, "
    "generating GST-compliant invoices, and maintaining customer credit ledgers. Existing "
    "tools either focus on accounting alone, lack modern user interfaces, or require "
    "expensive desktop installations that are unsuitable for mobile-first shopkeepers. "
    "This project presents the design and implementation of the AI Smart Inventory System, "
    "a full-stack web application that unifies inventory management, GST billing, customer "
    "credit (khata) tracking, and AI-driven business insights into a single platform.",

    "The system is built using the MERN stack — MongoDB, Express.js, React, and Node.js — "
    "with additional integrations for Google Gemini 2.5 Flash to power a conversational "
    "AI Copilot, Tesseract.js for optical character recognition of supplier invoices, and "
    "the pdfkit library for server-side generation of GST-compliant invoice PDFs. The "
    "application supports both English and Hindi user interfaces through the react-i18next "
    "library, addressing the linguistic diversity of Indian SMB users.",

    "Key technical contributions include: (i) a race-free, gap-free fiscal-year invoice "
    "numbering scheme implemented via an atomic counter collection in MongoDB; (ii) atomic "
    "stock deduction using conditional update operators that mathematically prevent oversell "
    "race conditions; (iii) an append-only customer ledger with a reversal-entry audit "
    "pattern that preserves financial integrity; and (iv) a tool-using AI Copilot that "
    "answers business queries by invoking typed helper functions on the live database.",

    "The system implements 47 RESTful API endpoints across 13 functional domains and 10 "
    "MongoDB collections. A daily cron job scans inventory for low-stock, out-of-stock, "
    "and dead-stock conditions, surfacing actionable alerts to the user. The user "
    "interface uses Tailwind CSS with a custom design system, supports a dark mode, and "
    "uses code-splitting to keep the initial JavaScript bundle below 300 kilobytes.",

    "The proposed system was tested against the requirements of a typical Indian SMB "
    "scenario, including intra-state and inter-state GST calculations, multi-line invoices "
    "with varying HSN codes, customer credit recording with running balance, and OCR-based "
    "goods receipt entry. The results demonstrate that the system successfully addresses "
    "the identified gaps in existing tools and provides a foundation for further "
    "extension into e-invoicing, multi-tenant SaaS deployment, and regional language coverage.",

    "Keywords: Inventory Management, GST Billing, MERN Stack, Artificial Intelligence, "
    "OCR, MongoDB, React, Node.js, Customer Ledger, SMB Software.",
]

for para in abstract_paras:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(para)
    r.font.size = Pt(12)

add_page_break(doc)

# ============ TABLE OF CONTENTS ============
add_centered(doc, "TABLE OF CONTENTS", size=22, bold=True, color=PRIMARY, space_before=24)
add_horizontal_rule(doc)
doc.add_paragraph()

toc_entries = [
    ("Certificate", "i"),
    ("Declaration", "ii"),
    ("Acknowledgement", "iii"),
    ("Abstract", "iv"),
    ("Table of Contents", "v"),
    ("List of Figures", "vi"),
    ("List of Tables", "vii"),
    ("List of Abbreviations", "viii"),
    ("", ""),
    ("Chapter 1: Introduction", "1"),
    ("    1.1 Overview", "1"),
    ("    1.2 Problem Statement", "2"),
    ("    1.3 Motivation", "3"),
    ("    1.4 Objectives of the Project", "4"),
    ("    1.5 Scope of the Project", "5"),
    ("    1.6 Organisation of the Report", "6"),
    ("", ""),
    ("Chapter 2: Literature Survey & Existing Systems", "7"),
    ("    2.1 Overview of Existing Systems", "7"),
    ("    2.2 Comparative Analysis", "10"),
    ("    2.3 Limitations of Existing Systems", "11"),
    ("    2.4 Proposed System", "12"),
    ("", ""),
    ("Chapter 3: System Analysis & Requirements", "13"),
    ("    3.1 Functional Requirements", "13"),
    ("    3.2 Non-Functional Requirements", "15"),
    ("    3.3 Hardware Requirements", "16"),
    ("    3.4 Software Requirements", "16"),
    ("    3.5 Feasibility Study", "17"),
    ("", ""),
    ("Chapter 4: System Design", "19"),
    ("    4.1 System Architecture", "19"),
    ("    4.2 Module Design", "21"),
    ("    4.3 Database Schema Design", "23"),
    ("    4.4 Entity-Relationship Diagram", "26"),
    ("    4.5 Data Flow Diagrams", "27"),
    ("    4.6 Use-Case Diagram", "29"),
    ("", ""),
    ("Chapter 5: Implementation", "30"),
    ("    5.1 Technology Stack", "30"),
    ("    5.2 Backend Implementation", "32"),
    ("    5.3 Frontend Implementation", "36"),
    ("    5.4 Algorithms & Key Logic", "39"),
    ("    5.5 AI Integration", "42"),
    ("    5.6 OCR Pipeline", "44"),
    ("    5.7 Security Implementation", "45"),
    ("", ""),
    ("Chapter 6: Testing", "47"),
    ("    6.1 Testing Strategy", "47"),
    ("    6.2 Unit Test Cases", "48"),
    ("    6.3 Integration Test Cases", "50"),
    ("    6.4 Test Results", "52"),
    ("", ""),
    ("Chapter 7: Results & Screenshots", "53"),
    ("", ""),
    ("Chapter 8: Conclusion", "60"),
    ("Chapter 9: Future Scope", "61"),
    ("References", "63"),
    ("Appendix A: Sample Code Listings", "65"),
    ("Appendix B: API Endpoint Reference", "68"),
]

for entry, page in toc_entries:
    if entry == "":
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    is_chapter = entry.startswith("Chapter") or entry in {
        "Certificate", "Declaration", "Acknowledgement", "Abstract", "Table of Contents",
        "List of Figures", "List of Tables", "List of Abbreviations",
        "References", "Appendix A: Sample Code Listings", "Appendix B: API Endpoint Reference",
    }
    pad_text = entry + " " + ("." * max(2, 90 - len(entry) - len(page)))
    r = p.add_run(pad_text)
    r.font.size = Pt(11)
    if is_chapter:
        r.font.bold = True
    r2 = p.add_run(" " + page)
    r2.font.size = Pt(11)
    r2.font.bold = is_chapter

add_page_break(doc)

# ============ LIST OF FIGURES ============
add_centered(doc, "LIST OF FIGURES", size=20, bold=True, color=PRIMARY, space_before=24)
add_horizontal_rule(doc)
doc.add_paragraph()
figs = [
    ("Figure 4.1", "High-Level System Architecture", "19"),
    ("Figure 4.2", "Three-Tier Application Layout", "20"),
    ("Figure 4.3", "Module Decomposition Diagram", "21"),
    ("Figure 4.4", "Entity-Relationship Diagram", "26"),
    ("Figure 4.5", "Level-0 Data Flow Diagram (Context Diagram)", "27"),
    ("Figure 4.6", "Level-1 Data Flow Diagram (Sales Module)", "28"),
    ("Figure 4.7", "Use-Case Diagram", "29"),
    ("Figure 5.1", "Technology Stack Visualisation", "31"),
    ("Figure 5.2", "Atomic Stock Deduction Sequence", "40"),
    ("Figure 5.3", "Invoice Number Allocation Flow", "41"),
    ("Figure 5.4", "AI Copilot Tool-Use Sequence", "43"),
    ("Figure 5.5", "OCR Goods-Receipt Pipeline", "44"),
    ("Figure 7.1", "Login Screen", "53"),
    ("Figure 7.2", "Dashboard with KPI Tiles", "54"),
    ("Figure 7.3", "Inventory Management Page", "55"),
    ("Figure 7.4", "Sales / GST Invoice Builder", "56"),
    ("Figure 7.5", "Generated GST Invoice PDF", "57"),
    ("Figure 7.6", "AI Insights Copilot Conversation", "58"),
    ("Figure 7.7", "OCR Scanner Review Screen", "59"),
]
add_table(doc, ["Figure", "Description", "Page"],
          [(f, d, p) for f, d, p in figs],
          col_widths_cm=[2.5, 11, 1.5])
add_page_break(doc)

# ============ LIST OF TABLES ============
add_centered(doc, "LIST OF TABLES", size=20, bold=True, color=PRIMARY, space_before=24)
add_horizontal_rule(doc)
doc.add_paragraph()
tabs = [
    ("Table 2.1", "Comparative Analysis of Existing Inventory Systems", "10"),
    ("Table 3.1", "Hardware Requirements", "16"),
    ("Table 3.2", "Software Requirements", "16"),
    ("Table 4.1", "Database Collections Overview", "23"),
    ("Table 4.2", "Product Schema Fields", "24"),
    ("Table 4.3", "Sale Schema Fields", "24"),
    ("Table 4.4", "Customer Schema Fields", "25"),
    ("Table 4.5", "KhataEntry Schema Fields", "25"),
    ("Table 5.1", "Backend Technology Versions", "30"),
    ("Table 5.2", "Frontend Technology Versions", "31"),
    ("Table 5.3", "API Endpoint Summary by Domain", "35"),
    ("Table 6.1", "Unit Test Cases", "48"),
    ("Table 6.2", "Integration Test Cases", "50"),
    ("Table 6.3", "Test Result Summary", "52"),
]
add_table(doc, ["Table", "Description", "Page"],
          [(t, d, p) for t, d, p in tabs],
          col_widths_cm=[2.5, 11, 1.5])
add_page_break(doc)

# ============ LIST OF ABBREVIATIONS ============
add_centered(doc, "LIST OF ABBREVIATIONS", size=20, bold=True, color=PRIMARY, space_before=24)
add_horizontal_rule(doc)
doc.add_paragraph()
abbrs = [
    ("API", "Application Programming Interface"),
    ("CGST", "Central Goods and Services Tax"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("DBMS", "Database Management System"),
    ("DFD", "Data Flow Diagram"),
    ("ER", "Entity-Relationship"),
    ("GRN", "Goods Received Note"),
    ("GST", "Goods and Services Tax"),
    ("GSTIN", "Goods and Services Tax Identification Number"),
    ("HSN", "Harmonised System of Nomenclature"),
    ("HTTP", "Hypertext Transfer Protocol"),
    ("IGST", "Integrated Goods and Services Tax"),
    ("JSON", "JavaScript Object Notation"),
    ("JWT", "JSON Web Token"),
    ("LLM", "Large Language Model"),
    ("MERN", "MongoDB, Express.js, React, Node.js"),
    ("OCR", "Optical Character Recognition"),
    ("PDF", "Portable Document Format"),
    ("RBAC", "Role-Based Access Control"),
    ("REST", "Representational State Transfer"),
    ("SaaS", "Software as a Service"),
    ("SGST", "State Goods and Services Tax"),
    ("SKU", "Stock Keeping Unit"),
    ("SMB", "Small and Medium Business"),
    ("SPA", "Single Page Application"),
    ("UI", "User Interface"),
    ("UPI", "Unified Payments Interface"),
    ("URL", "Uniform Resource Locator"),
    ("UX", "User Experience"),
    ("XML", "Extensible Markup Language"),
]
add_table(doc, ["Abbreviation", "Expansion"], abbrs, col_widths_cm=[3, 12])
add_page_break(doc)

# ============================================================
# CHAPTER 1 — INTRODUCTION
# ============================================================
add_chapter_heading(doc, 1, "Introduction")

add_section_heading(doc, "1.1", "Overview")
add_para(doc,
    "Inventory management and billing form the operational backbone of every retail "
    "business, from a single-shop kirana to a multi-location distributor. In the Indian "
    "context, this responsibility is significantly amplified by the requirements of the "
    "Goods and Services Tax (GST) regime, which mandates structured invoicing with "
    "Harmonised System of Nomenclature (HSN) codes, state-wise tax splits between Central "
    "GST (CGST), State GST (SGST), and Integrated GST (IGST), and the maintenance of "
    "complete audit-ready records. Alongside these compliance pressures, small and medium "
    "businesses (SMBs) also extend significant credit to their customers through informal "
    "ledger systems known locally as khata, where balances are tracked in a paper notebook "
    "rather than in formal accounting software.")
add_para(doc,
    "The AI Smart Inventory System is a comprehensive web-based platform that unifies "
    "these traditionally separate concerns — product cataloguing, stock movement tracking, "
    "GST-compliant invoicing, customer credit ledger management, supplier records, and "
    "operational analytics — into a single, modern, mobile-friendly application. The "
    "system additionally introduces two capabilities that are largely absent from existing "
    "Indian SMB tools: an Optical Character Recognition (OCR) pipeline that converts a "
    "photograph of a supplier's invoice into a structured Goods Received Note in seconds, "
    "and an Artificial Intelligence Copilot powered by Google Gemini 2.5 Flash that "
    "answers natural-language business questions by querying the live application database "
    "through typed tool functions.")
add_para(doc,
    "This report documents the analysis, design, implementation, and testing of the AI "
    "Smart Inventory System. The application is built on the MERN technology stack — "
    "MongoDB as the document database, Express.js as the web framework, React 19 as the "
    "single-page application library, and Node.js 20 as the JavaScript runtime — and "
    "augments this stack with Tesseract.js for OCR, the pdfkit library for server-side "
    "PDF generation, react-i18next for English/Hindi internationalisation, and Tailwind "
    "CSS for the user-interface design system.")

add_section_heading(doc, "1.2", "Problem Statement")
add_para(doc,
    "Despite the digitisation push of the Goods and Services Tax regime, a significant "
    "fraction of Indian SMBs continue to operate using a fragmented combination of paper "
    "bill books, spreadsheets, locally installed accounting software (typically Tally), "
    "and credit-tracking applications such as Khatabook. Each of these tools addresses "
    "one part of the problem but leaves substantial workflow gaps that the shopkeeper "
    "has to bridge manually.")
add_para(doc,
    "The specific problems addressed by this project are:")
add_bullet(doc, "Existing accounting software is desktop-bound and visually outdated, "
                "making it difficult to use on mobile devices that have become the primary "
                "computing platform for many small-business owners.")
add_bullet(doc, "Customer credit (khata) management and GST invoicing are rarely combined "
                "in a single tool. Users either choose Khatabook for credit tracking and "
                "lose GST features, or choose Tally/Vyapar for GST and lose the simplicity "
                "of a credit notebook.")
add_bullet(doc, "Manual entry of supplier invoices into the system is slow and error-prone. "
                "A medium-sized distributor receiving twenty supplier invoices per week "
                "spends hours typing them into the inventory system.")
add_bullet(doc, "There is no readily available AI assistant that can answer ad-hoc business "
                "questions such as \"which products are about to run out?\" or \"who owes "
                "me the most money this month?\" by reading the user's own data.")
add_bullet(doc, "Hindi and other regional language support in existing SMB software is "
                "often a partial translation rather than a first-class language, which "
                "creates a usability barrier for shopkeepers in Tier-2 and Tier-3 cities.")

add_section_heading(doc, "1.3", "Motivation")
add_para(doc,
    "The motivation for this project arises from three converging trends that have made "
    "an integrated, AI-powered SMB application both technically feasible and commercially "
    "relevant:")
add_para(doc,
    "First, the cost of inference for capable Large Language Models has fallen dramatically "
    "with the release of efficient models such as Google Gemini 2.5 Flash. It is now "
    "economically practical to embed a tool-using AI assistant into an SMB application "
    "without making the application prohibitively expensive to operate.")
add_para(doc,
    "Second, the Unified Payments Interface (UPI) has reached near-universal adoption in "
    "Indian retail, making it natural to embed Scan-to-Pay QR codes directly into "
    "every printed invoice. This was simply not possible at scale even five years ago.")
add_para(doc,
    "Third, the GST regime continues to evolve toward greater digitisation through "
    "mechanisms such as e-invoicing (with Invoice Reference Numbers and signed QR codes), "
    "and the turnover threshold for mandatory e-invoicing has been progressively lowered. "
    "Indian SMBs above the threshold need a software path that can grow into these "
    "advanced compliance requirements without forcing a tool migration every two years.")

add_section_heading(doc, "1.4", "Objectives of the Project")
add_para(doc, "The primary objectives of the AI Smart Inventory System are as follows:")
for obj in [
    "To design and implement a unified web-based platform that integrates inventory "
    "management, GST-compliant invoicing, customer credit (khata) ledger, and supplier "
    "management into a single coherent application.",
    "To implement a robust GST calculation engine that correctly handles intra-state "
    "transactions (with CGST and SGST splits) and inter-state transactions (with IGST), "
    "based on the seller's and buyer's registered states.",
    "To develop a race-free, gap-free invoice numbering scheme using atomic database "
    "operations, ensuring that no two concurrent transactions can ever produce the same "
    "or skipped invoice number.",
    "To build an Optical Character Recognition (OCR) pipeline that converts supplier "
    "invoice images into structured Goods Received Notes, with human-in-the-loop "
    "verification before stock-in is committed.",
    "To integrate an Artificial Intelligence Copilot that uses the Google Gemini "
    "2.5 Flash model with typed tool functions to answer natural-language business "
    "queries based on live application data.",
    "To implement an automated daily background job that scans inventory for low-stock, "
    "out-of-stock, and dead-stock conditions and surfaces actionable alerts to the user.",
    "To provide a multilingual user interface supporting English and Hindi, designed "
    "with mobile-first responsiveness, accessibility, and a dark mode option.",
    "To enforce data integrity through atomic operations, append-only ledger entries, "
    "schema-level validation, and role-based access control for multi-user shop deployments.",
]:
    add_bullet(doc, obj)

add_section_heading(doc, "1.5", "Scope of the Project")
add_para(doc,
    "The scope of this project covers the complete design, implementation, and "
    "demonstration of a working web application that addresses the workflow needs of "
    "a typical Indian retail SMB. The following modules and capabilities are within scope:")
for s in [
    "User authentication, registration, and role-based access control (admin, manager, staff)",
    "Product catalogue with SKU, HSN code, GST rate, cost and selling price, reorder threshold, and supplier linkage",
    "Sales transaction recording with multi-line items, automatic GST split, and PDF invoice generation",
    "Customer master with phone, GSTIN, address, state, and credit limit",
    "Append-only customer ledger (khata) with payments, adjustments, and reversal entries",
    "Supplier directory with linked products and stock-IN history",
    "OCR-based supplier invoice scanning with review-and-confirm workflow",
    "AI Copilot with five live tool functions for inventory and sales queries",
    "Daily smart-alerts cron for low-stock, out-of-stock, and dead-stock detection",
    "Dashboard with KPI tiles and analytics page with charts",
    "English and Hindi user-interface localisation with persistent language preference",
    "Settings page for workspace configuration and notification preferences",
]:
    add_bullet(doc, s)

add_para(doc, "The following are explicitly out of the scope of this project:")
for s in [
    "GST e-invoicing (Invoice Reference Number generation through a registered GSP)",
    "GSTR-1 and GSTR-3B JSON export for direct portal upload",
    "Multi-tenant SaaS deployment with subscription billing",
    "Native mobile applications for Android and iOS (web-responsive only)",
    "Hardware barcode scanner integration (camera-based scanning is in scope as an extension)",
    "Direct integration with marketplaces such as Shopify, Amazon, or Meesho",
]:
    add_bullet(doc, s)

add_section_heading(doc, "1.6", "Organisation of the Report")
add_para(doc,
    "The remainder of this report is organised as follows. Chapter 2 surveys existing "
    "inventory and billing software in the Indian SMB market and analyses their "
    "limitations. Chapter 3 captures the functional and non-functional requirements "
    "and presents a feasibility study. Chapter 4 describes the system design — "
    "architecture, modules, database schema, ER diagram, and data flow diagrams. "
    "Chapter 5 covers implementation details across the backend, frontend, AI "
    "integration, OCR pipeline, and security. Chapter 6 documents the testing "
    "strategy and results. Chapter 7 presents output screenshots. Chapter 8 concludes "
    "the report and Chapter 9 outlines the future scope. The references and "
    "appendices follow.")

add_page_break(doc)

# ============================================================
# CHAPTER 2 — LITERATURE SURVEY
# ============================================================
add_chapter_heading(doc, 2, "Literature Survey & Existing Systems")

add_section_heading(doc, "2.1", "Overview of Existing Systems")
add_para(doc,
    "Several software products currently serve the Indian SMB inventory and billing "
    "market. Each has distinct strengths, weaknesses, and a distinct user base. A "
    "review of these systems was conducted before the design of the proposed system "
    "to identify the workflow gaps that this project would attempt to close.")

add_subsection_heading(doc, "2.1.1", "Tally Prime")
add_para(doc,
    "Tally Prime is a long-standing accounting and inventory management software that "
    "is widely used by chartered accountants and mid-sized businesses across India. "
    "It offers comprehensive double-entry bookkeeping, GST return preparation, "
    "inventory tracking with stock groups and categories, and a robust voucher-based "
    "transaction model.")
add_para(doc,
    "Tally's strengths are its accountant-grade financial controls, its acceptance as "
    "a de facto standard by chartered accountants, and its mature handling of "
    "complex GST scenarios. Its primary limitations are a desktop-only deployment "
    "model that does not run natively on smartphones, a user interface that has "
    "not been substantially modernised in the past decade, and a steep learning "
    "curve that requires familiarity with accounting concepts.")

add_subsection_heading(doc, "2.1.2", "Vyapar")
add_para(doc,
    "Vyapar is a popular billing and inventory application aimed specifically at "
    "small Indian businesses. It runs on Android, iOS, and Windows, supports multiple "
    "Indian languages, and offers GST invoicing, expense tracking, basic inventory "
    "management, and customer payment reminders.")
add_para(doc,
    "Vyapar's strengths are its accessible price point, broad language support, and "
    "mobile-first interface. Its limitations include the absence of any genuine AI "
    "capabilities, limited support for OCR-based supplier invoice ingestion, and a "
    "relatively basic credit-ledger model compared to dedicated khata applications.")

add_subsection_heading(doc, "2.1.3", "Khatabook")
add_para(doc,
    "Khatabook is a digital cashbook and credit-ledger application designed to "
    "replace the traditional paper khata notebook. It supports recording credit "
    "transactions, sending payment reminders, and maintaining customer balances.")
add_para(doc,
    "Khatabook's strength is the simplicity of its credit-tracking model, which maps "
    "directly to the mental model of an Indian shopkeeper. Its limitations are the "
    "lack of GST invoice generation, the absence of inventory management, and the "
    "limited reporting and analytics capabilities.")

add_subsection_heading(doc, "2.1.4", "Zoho Books")
add_para(doc,
    "Zoho Books is a cloud-based accounting and invoicing platform from Zoho "
    "Corporation. It supports GST-compliant invoicing, automated payment reminders, "
    "bank reconciliation, and integration with the broader Zoho product suite.")
add_para(doc,
    "Zoho Books offers a modern web-based user interface and broad feature coverage. "
    "Its limitations for the Indian SMB segment are an English-first design that does "
    "not translate naturally to Hindi-speaking users, a price point above the level "
    "that small kirana operators are willing to pay, and the absence of a dedicated "
    "khata-style credit ledger workflow.")

add_section_heading(doc, "2.2", "Comparative Analysis")
add_para(doc,
    "The following table compares the four reviewed systems against the proposed AI "
    "Smart Inventory System on the dimensions most relevant to an Indian SMB user.")

add_table(doc,
    ["Capability", "Proposed", "Tally", "Vyapar", "Khatabook", "Zoho Books"],
    [
        ["GST invoicing (CGST/SGST/IGST)", "Yes", "Yes", "Yes", "No", "Yes"],
        ["Customer khata ledger", "Yes", "Manual", "Basic", "Yes", "Partial"],
        ["OCR supplier invoice ingest", "Yes", "No", "No", "No", "No"],
        ["AI Copilot with live tool-use", "Yes", "No", "No", "No", "Basic"],
        ["Hindi UI (first-class)", "Yes", "Partial", "Yes", "Yes", "No"],
        ["Web-based (any device)", "Yes", "No", "Partial", "Yes", "Yes"],
        ["Open architecture for extension", "Yes", "No", "No", "No", "Limited"],
        ["Modern dark-mode UX", "Yes", "No", "No", "Partial", "Partial"],
    ],
    col_widths_cm=[5, 1.7, 1.7, 1.7, 1.7, 1.7])

add_section_heading(doc, "2.3", "Limitations of Existing Systems")
add_para(doc,
    "From the comparative review, the following limitations of the existing landscape "
    "emerge as concrete opportunities for the proposed system:")
for lim in [
    "No single existing tool combines GST invoicing, khata credit tracking, supplier OCR, and AI insights in one application.",
    "OCR-driven goods receipt entry is largely unimplemented in this market segment.",
    "AI assistance is either entirely absent or limited to canned report templates rather than live, tool-using conversational queries.",
    "Hindi support is uneven, often missing from the invoice PDF itself even when present in the application UI.",
    "Mobile-first responsive design is rare in the established players, who continue to optimise for desktop deployments.",
]:
    add_bullet(doc, lim)

add_section_heading(doc, "2.4", "Proposed System")
add_para(doc,
    "The proposed AI Smart Inventory System is designed to address each of the "
    "limitations identified above. Its key differentiators are:")
for d in [
    "A single application covering inventory, billing, khata, and supplier management without context switching.",
    "Server-side OCR ingestion of supplier invoices via Tesseract.js with a mandatory human review step to ensure data integrity.",
    "A conversational AI Copilot built on Google Gemini 2.5 Flash with five typed tool functions that operate on live application data.",
    "First-class Hindi user-interface support driven by react-i18next, including localisation of the invoice PDF.",
    "A mobile-first web architecture using React 19, Tailwind CSS, and code-splitting that runs equally well on a smartphone browser and a desktop browser.",
    "Built-in dark mode and a coherent design system using a Bahi-Red and Paper-Ivory palette inspired by traditional Indian accountant ledger books.",
]:
    add_bullet(doc, d)

add_page_break(doc)

# ============================================================
# CHAPTER 3 — SYSTEM ANALYSIS & REQUIREMENTS
# ============================================================
add_chapter_heading(doc, 3, "System Analysis & Requirements")

add_section_heading(doc, "3.1", "Functional Requirements")
add_para(doc,
    "The functional requirements of the system specify the behaviour that the application "
    "must exhibit to support the daily workflow of an Indian SMB user. They are organised "
    "below by functional module.")

add_subsection_heading(doc, "3.1.1", "Authentication & User Management")
for r in [
    "FR-AUTH-1: The system shall allow new users to register with name, email, and password.",
    "FR-AUTH-2: The system shall authenticate users via email and password, returning a JSON Web Token on success.",
    "FR-AUTH-3: The system shall support three user roles: admin, manager, and staff, each with progressively reduced privileges.",
    "FR-AUTH-4: The system shall allow admin users to invite, edit, and deactivate other users in the same workspace.",
    "FR-AUTH-5: The system shall allow users to update their profile and change password.",
]:
    add_bullet(doc, r)

add_subsection_heading(doc, "3.1.2", "Inventory Management")
for r in [
    "FR-INV-1: The system shall allow creation, reading, updating, and soft-deletion of products.",
    "FR-INV-2: Each product shall capture name, SKU (unique), category, HSN code, GST rate, cost price, selling price, current stock, low-stock threshold, supplier, and optional barcode.",
    "FR-INV-3: The system shall support stock adjustments classified as IN, OUT, or ADJUST with a mandatory reason field.",
    "FR-INV-4: The system shall expose a barcode-lookup endpoint for scan-to-sell flows.",
    "FR-INV-5: The system shall expose a low-stock query endpoint that lists products below their threshold.",
]:
    add_bullet(doc, r)

add_subsection_heading(doc, "3.1.3", "Sales & GST Invoicing")
for r in [
    "FR-SALE-1: The system shall allow creation of multi-line sales transactions with automatic GST calculation.",
    "FR-SALE-2: GST shall be split into CGST and SGST for intra-state sales, and into IGST for inter-state sales, based on seller and buyer state codes.",
    "FR-SALE-3: Invoice numbers shall follow the format INV-{fiscal_year}-{5-digit-sequence} and shall be allocated atomically with no gaps and no duplicates.",
    "FR-SALE-4: The HSN code on each line item shall be frozen at sale time so historical invoices remain accurate even if the underlying product is later edited.",
    "FR-SALE-5: The system shall generate a server-side A4 PDF invoice with seller details, buyer details, GST split table, amount in words, and an embedded UPI Scan-to-Pay QR code.",
    "FR-SALE-6: The system shall provide a WhatsApp share link with prefilled invoice number and amount.",
    "FR-SALE-7: Stock for sold items shall be deducted atomically with a conditional update that prevents oversell.",
]:
    add_bullet(doc, r)

add_subsection_heading(doc, "3.1.4", "Customers & Khata Ledger")
for r in [
    "FR-CUST-1: The system shall allow creation, reading, updating, and soft-deletion of customer records.",
    "FR-CUST-2: Each customer shall capture name, phone (Indian format with +91 prefix), email (optional), GSTIN (optional, validated to 15 characters), address, Indian state, opening balance, and credit limit.",
    "FR-CUST-3: The system shall maintain an append-only ledger (KhataEntry) for each customer.",
    "FR-CUST-4: A sale on credit shall automatically post a debit entry to the customer's ledger.",
    "FR-CUST-5: A payment from the customer shall post a credit entry with mode (cash, UPI, cheque, bank), receipt number, and reference identifier.",
    "FR-CUST-6: A reversal of any prior entry shall create a new entry with reversal-of reference and shall mark both entries as reversed; the original entry shall not be mutated.",
    "FR-CUST-7: The system shall expose a top-debtors query and a per-customer statement (PDF or JSON) endpoint.",
]:
    add_bullet(doc, r)

add_subsection_heading(doc, "3.1.5", "OCR Supplier Invoice")
for r in [
    "FR-OCR-1: The system shall accept upload of a supplier invoice image in JPG or PNG format up to 10 MB.",
    "FR-OCR-2: The system shall extract text from the uploaded image using Tesseract.js.",
    "FR-OCR-3: The system shall parse the extracted text to identify invoice number, vendor name, line items, and totals using regular-expression heuristics.",
    "FR-OCR-4: The system shall present the parsed data to the user for review and correction before any database mutation.",
    "FR-OCR-5: On user confirmation, the system shall create a Goods Received Note, atomically increment stock for each line item, and write IN-type Transaction audit rows linked to the original image.",
]:
    add_bullet(doc, r)

add_subsection_heading(doc, "3.1.6", "AI Insights Copilot")
for r in [
    "FR-AI-1: The system shall provide a chat interface that streams responses from the Google Gemini 2.5 Flash model.",
    "FR-AI-2: The Copilot shall expose at least five typed tool functions: get_low_stock, get_top_movers, get_dead_stock, get_gst_summary, and get_supplier_list.",
    "FR-AI-3: The Copilot shall display suggested-question chips grouped by intent (inventory, sales, suppliers, dead-stock, operations).",
    "FR-AI-4: AI chat requests shall be rate-limited to 20 requests per minute per IP address.",
    "FR-AI-5: The system shall provide additional AI-driven endpoints for 30-day demand forecast, dead-stock list, per-product reorder suggestion, and sales-velocity trends.",
]:
    add_bullet(doc, r)

add_subsection_heading(doc, "3.1.7", "Smart Alerts")
for r in [
    "FR-ALERT-1: A daily cron job shall scan all products at a configurable time (default 9:00 IST) for OUT_OF_STOCK, LOW_STOCK, and DEAD_STOCK conditions.",
    "FR-ALERT-2: Alerts shall be deduplicated so that a recurring condition does not generate multiple identical alerts.",
    "FR-ALERT-3: An admin user shall be able to trigger the alert scan on demand.",
    "FR-ALERT-4: Active alert count shall be visible as a badge in the top navigation bell icon.",
]:
    add_bullet(doc, r)

add_section_heading(doc, "3.2", "Non-Functional Requirements")

add_subsection_heading(doc, "3.2.1", "Performance")
for r in [
    "NFR-PERF-1: The system shall serve API responses in under 200 ms for cached read paths under nominal load.",
    "NFR-PERF-2: The initial JavaScript bundle of the frontend application shall be under 300 kB after gzip compression.",
    "NFR-PERF-3: PDF invoice generation shall complete in under 1.5 seconds for invoices with up to 50 line items.",
]:
    add_bullet(doc, r)

add_subsection_heading(doc, "3.2.2", "Security")
for r in [
    "NFR-SEC-1: User passwords shall be stored using bcryptjs hashing with a minimum of 10 rounds.",
    "NFR-SEC-2: All authenticated endpoints shall require a valid JSON Web Token in the Authorization header.",
    "NFR-SEC-3: Authentication endpoints shall be rate-limited to 15 requests per minute per IP address.",
    "NFR-SEC-4: Input on every POST and PUT route shall be validated against a Zod schema; invalid input shall return a 400 status with field-level error messages.",
]:
    add_bullet(doc, r)

add_subsection_heading(doc, "3.2.3", "Reliability")
for r in [
    "NFR-REL-1: Invoice numbers shall be allocated using an atomic findOneAndUpdate operation on a Counter collection, guaranteeing no duplicates and no gaps.",
    "NFR-REL-2: Stock deduction shall use a conditional atomic update that prevents oversell under concurrent transactions.",
    "NFR-REL-3: KhataEntry records shall be append-only; reversals shall be implemented as new entries that reference the original.",
]:
    add_bullet(doc, r)

add_subsection_heading(doc, "3.2.4", "Usability")
for r in [
    "NFR-USE-1: The user interface shall be responsive and shall function on screens from 360 px to 1920 px wide.",
    "NFR-USE-2: The application shall support both English and Hindi languages with a persistent per-user preference.",
    "NFR-USE-3: The application shall support a dark mode in addition to the default light mode.",
    "NFR-USE-4: All interactive controls shall have a minimum touch target of 44 by 44 pixels for mobile usability.",
]:
    add_bullet(doc, r)

add_section_heading(doc, "3.3", "Hardware Requirements")
add_table(doc, ["Component", "Minimum", "Recommended"],
    [
        ["Processor", "Intel Core i3 / AMD Ryzen 3", "Intel Core i5 / AMD Ryzen 5 or higher"],
        ["RAM", "4 GB", "8 GB or higher"],
        ["Storage", "20 GB free", "50 GB SSD"],
        ["Display", "13 inch / 1366 x 768", "15 inch / 1920 x 1080"],
        ["Internet", "1 Mbps for development", "10 Mbps for cloud deployment"],
        ["Server (production)", "2 vCPU / 4 GB RAM cloud VM", "4 vCPU / 8 GB RAM with autoscaling"],
        ["Mobile (end user)", "Any Android 9+ / iOS 14+ smartphone with 2 GB RAM", "Android 11+ / iOS 15+, 4 GB RAM"],
    ],
    col_widths_cm=[4, 5.5, 5.5])

add_section_heading(doc, "3.4", "Software Requirements")
add_table(doc, ["Layer", "Technology", "Version"],
    [
        ["Operating System (development)", "Windows 10 / 11, macOS 12+, or Linux", "Any current"],
        ["Database", "MongoDB (local or Atlas)", "6.x or higher"],
        ["JavaScript runtime", "Node.js", "20.x LTS"],
        ["Web framework (backend)", "Express.js", "5.x"],
        ["ODM", "Mongoose", "9.x"],
        ["Frontend library", "React", "19.x"],
        ["Build tool", "Vite", "8.x"],
        ["CSS framework", "Tailwind CSS", "3.x"],
        ["AI SDK", "@google/generative-ai", "0.24+"],
        ["OCR engine", "Tesseract.js", "7.x"],
        ["PDF library", "pdfkit", "0.18+"],
        ["i18n", "react-i18next", "Current"],
        ["Browser (end user)", "Chrome / Edge / Firefox / Safari", "Last 2 major versions"],
    ],
    col_widths_cm=[5, 6, 4])

add_section_heading(doc, "3.5", "Feasibility Study")

add_subsection_heading(doc, "3.5.1", "Technical Feasibility")
add_para(doc,
    "All technologies required for the project are mature, well-documented, and "
    "freely available under permissive open-source licenses, with the exception of "
    "the Google Gemini API which operates on a metered pay-per-use model with a "
    "generous free tier sufficient for development and demonstration. The project "
    "does not require any specialised hardware or proprietary infrastructure. The "
    "MERN stack runs on commodity cloud virtual machines and on most laptops used "
    "for academic project work. The project is therefore technically feasible.")

add_subsection_heading(doc, "3.5.2", "Operational Feasibility")
add_para(doc,
    "The application is designed around the existing daily workflow of Indian SMB "
    "shopkeepers — recording sales, managing stock, accepting payments, tracking "
    "credit, and reordering from suppliers. The user interface uses familiar "
    "metaphors (the khata book, the bahi ledger, the invoice register) and the "
    "Hindi-language option lowers the linguistic barrier. The system can be operated "
    "from any modern web browser, making it accessible from desktop, laptop, "
    "tablet, and smartphone. The project is therefore operationally feasible.")

add_subsection_heading(doc, "3.5.3", "Economic Feasibility")
add_para(doc,
    "The project relies almost entirely on free and open-source software. The only "
    "operational cost is cloud hosting for production deployment (estimated at less "
    "than ten dollars per month for a small instance) and the AI API usage cost "
    "(controlled by per-IP rate limits and bounded conversation context). These "
    "costs are negligible at academic project scale. The project is therefore "
    "economically feasible.")

add_page_break(doc)

# ============================================================
# CHAPTER 4 — SYSTEM DESIGN
# ============================================================
add_chapter_heading(doc, 4, "System Design")

add_section_heading(doc, "4.1", "System Architecture")
add_para(doc,
    "The AI Smart Inventory System follows a classic three-tier web application "
    "architecture, augmented with an external AI service tier and a scheduled "
    "background-job tier. The presentation tier is a React-based single-page "
    "application served by the Vite development server (in development) or "
    "as static assets (in production). The application tier is an Express.js "
    "REST API running on Node.js, which encapsulates all business logic, "
    "validation, and integrations. The data tier is MongoDB, accessed through "
    "the Mongoose Object-Document Mapper.")

add_placeholder_box(doc, "Figure 4.1 — High-Level System Architecture (insert architecture diagram here)")

add_para(doc,
    "External integrations include the Google Gemini API for the AI Copilot and "
    "Tesseract.js (executed on the server) for OCR. A node-cron based scheduler "
    "runs a daily background job at 09:00 IST to scan inventory and generate "
    "smart alerts.")

add_placeholder_box(doc, "Figure 4.2 — Three-Tier Application Layout")

add_section_heading(doc, "4.2", "Module Design")
add_para(doc,
    "The system is decomposed into thirteen functional modules, each owning a "
    "specific aspect of the application's behaviour. Each module groups its routes, "
    "validators, controller, and service code under a common naming convention.")

add_placeholder_box(doc, "Figure 4.3 — Module Decomposition Diagram")

add_table(doc,
    ["Module", "Primary Responsibility"],
    [
        ["Auth", "Registration, login, JWT issuance, role enforcement"],
        ["Product", "Product catalogue CRUD and stock adjustments"],
        ["Sale", "Sales transactions, GST split, invoice numbering, PDF generation"],
        ["Customer", "Customer master and outstanding-balance cache"],
        ["Khata", "Append-only ledger, payments, adjustments, reversals, statements"],
        ["Supplier", "Supplier directory and supplier-product linkage"],
        ["Transaction", "Stock movement audit log"],
        ["Alert", "Smart alert listing, dismissal, on-demand trigger"],
        ["Analytics", "Dashboard KPIs and reporting"],
        ["OCR", "Supplier invoice upload, extraction, save"],
        ["AI", "Gemini Copilot chat, predict, insights, dead-stock, reorder, trends"],
        ["Settings", "Workspace configuration, AI config, notifications, password"],
        ["Health", "Liveness probe and version endpoint"],
    ],
    col_widths_cm=[4, 11])

add_section_heading(doc, "4.3", "Database Schema Design")
add_para(doc,
    "The system uses MongoDB as the primary data store. The schema is organised "
    "into ten collections. The following table summarises each collection and its "
    "purpose. Detailed field listings for the most important schemas follow.")

add_table(doc,
    ["Collection", "Purpose"],
    [
        ["User", "Authentication credentials and role"],
        ["Product", "Product catalogue with stock and pricing"],
        ["Sale", "Sales transactions and embedded line items"],
        ["Customer", "Customer master with outstanding-balance cache"],
        ["KhataEntry", "Append-only ledger lines per customer"],
        ["Supplier", "Supplier directory"],
        ["Transaction", "Stock movement audit log"],
        ["Settings", "Per-user workspace configuration"],
        ["Alert", "Inventory alerts generated by the cron"],
        ["Counter", "Atomic sequence allocator for invoice numbers"],
    ],
    col_widths_cm=[4, 11])

add_subsection_heading(doc, "4.3.1", "Product Schema")
add_table(doc, ["Field", "Type", "Constraints", "Description"],
    [
        ["sku", "String", "Unique, required", "Stock-keeping unit identifier"],
        ["name", "String", "Required", "Product display name"],
        ["category", "String", "Optional", "Category for grouping"],
        ["price", "Number", "Required, >= 0", "Selling price in INR"],
        ["costPrice", "Number", "Required, >= 0", "Cost / purchase price in INR"],
        ["hsnCode", "String", "4-8 digits", "Harmonised System of Nomenclature code"],
        ["barcode", "String", "Unique, sparse", "Optional barcode for scan-to-sell"],
        ["stock", "Number", "Required, >= 0", "Current stock quantity"],
        ["lowStockThreshold", "Number", "Default 5", "Low-stock alert threshold"],
        ["supplierId", "ObjectId", "Ref: Supplier", "Linked supplier"],
        ["isDeleted", "Boolean", "Default false", "Soft-delete flag"],
    ],
    col_widths_cm=[3.5, 2.5, 3.5, 5.5])

add_subsection_heading(doc, "4.3.2", "Sale Schema")
add_table(doc, ["Field", "Type", "Description"],
    [
        ["invoiceNumber", "String (unique)", "Format INV-{year}-{5-digit}"],
        ["customer", "Embedded object", "Snapshot of customer at sale time"],
        ["seller", "Embedded object", "Snapshot of seller workspace at sale time"],
        ["items", "Array", "Line items (productId, name, qty, unitPrice, hsnCode, taxes)"],
        ["subtotal", "Number", "Sum of (qty x unitPrice) before tax"],
        ["discount", "Number", "Discount applied (must be <= subtotal)"],
        ["gst", "Object", "{ isInterstate, cgstRate, sgstRate, igstRate, cgstAmt, sgstAmt, igstAmt }"],
        ["taxAmount", "Number", "Total tax (cgst + sgst + igst)"],
        ["total", "Number", "Final invoice total in INR"],
        ["paymentMode", "String enum", "cash | upi | bank | cheque | card | credit"],
        ["customerId", "ObjectId", "Ref: Customer"],
        ["createdBy", "ObjectId", "Ref: User"],
    ],
    col_widths_cm=[3.5, 3, 8.5])

add_subsection_heading(doc, "4.3.3", "Customer Schema")
add_table(doc, ["Field", "Type", "Description"],
    [
        ["name", "String", "Customer display name"],
        ["phone", "String", "Indian format: +91 followed by 10 digits starting with 6-9"],
        ["email", "String", "Optional"],
        ["gstin", "String", "Optional, 15 alphanumeric characters when present"],
        ["address", "String", "Free text address"],
        ["state", "String", "One of 36 Indian states/UTs (drives GST split)"],
        ["openingBalance", "Number", "Starting balance carried forward"],
        ["creditLimit", "Number", "Soft warning threshold; not enforced as a hard block"],
        ["outstandingBalance", "Number", "Denormalised cache, kept in sync with KhataEntry"],
        ["lastTransactionAt", "Date", "Timestamp of most recent ledger activity"],
        ["isActive", "Boolean", "Active flag (soft delete)"],
    ],
    col_widths_cm=[3.5, 2.5, 9])

add_subsection_heading(doc, "4.3.4", "KhataEntry Schema")
add_table(doc, ["Field", "Type", "Description"],
    [
        ["customerId", "ObjectId", "Ref: Customer"],
        ["voucherType", "Enum", "Sale | Payment | Refund | Adjustment | OpeningBalance"],
        ["direction", "Enum", "debit | credit"],
        ["amount", "Number", "Entry amount in INR"],
        ["runningBalance", "Number", "Frozen balance after this entry"],
        ["entryDate", "Date", "Logical entry date (may differ from createdAt)"],
        ["mode", "Enum", "cash | upi | bank | cheque | card | (none for non-payment)"],
        ["receiptNumber", "String", "RCPT-YYYY-NNNNN for payments and refunds"],
        ["chequeNumber", "String", "Optional, when mode = cheque"],
        ["upiTxnId", "String", "Optional, when mode = upi"],
        ["reversalOf", "ObjectId", "Self-reference when this entry reverses another"],
        ["isReversed", "Boolean", "True for both entries in a reversal pair"],
    ],
    col_widths_cm=[3.5, 2.5, 9])

add_section_heading(doc, "4.4", "Entity-Relationship Diagram")
add_placeholder_box(doc, "Figure 4.4 — Entity-Relationship Diagram (insert ER diagram)")

add_para(doc,
    "The principal entity relationships in the system are:")
for r in [
    "User : Settings = 1 : 1",
    "User : Customer = 1 : N",
    "User : Sale = 1 : N (createdBy)",
    "Customer : KhataEntry = 1 : N",
    "Customer : Sale = 1 : N",
    "Sale : Product = N : M (through embedded items array)",
    "Product : Transaction = 1 : N",
    "Product : Alert = 1 : N (cron-generated)",
    "Supplier : Product = 1 : N",
    "Counter : (no relation, used as a sequence allocator)",
]:
    add_bullet(doc, r)

add_section_heading(doc, "4.5", "Data Flow Diagrams")

add_subsection_heading(doc, "4.5.1", "Level-0 (Context Diagram)")
add_para(doc,
    "The Level-0 Data Flow Diagram captures the system as a single process and "
    "identifies its external entities — the SMB User (admin / manager / staff), "
    "the External AI service (Google Gemini), and the persisted MongoDB datastore. "
    "User actions and authentication credentials flow into the system; invoice "
    "PDFs, alerts, AI responses, and reports flow back to the user.")
add_placeholder_box(doc, "Figure 4.5 — Level-0 Data Flow Diagram (Context Diagram)")

add_subsection_heading(doc, "4.5.2", "Level-1 (Sales Module)")
add_para(doc,
    "The Level-1 Data Flow Diagram for the Sales module decomposes the create-sale "
    "operation into its constituent processes: validate input, allocate invoice "
    "number from Counter, deduct stock atomically, calculate GST split, post khata "
    "entry on credit sales, persist Sale document, and return the invoice "
    "identifier to the user. The diagram clearly shows the dependency on the "
    "Counter, Product, and KhataEntry stores within a single logical transaction.")
add_placeholder_box(doc, "Figure 4.6 — Level-1 Data Flow Diagram (Sales Module)")

add_section_heading(doc, "4.6", "Use-Case Diagram")
add_para(doc,
    "The use-case diagram identifies the actors and their interactions with the "
    "system. The primary actors are the Admin, Manager, and Staff roles. The "
    "secondary actor is the AI service. Key use cases include: register/login, "
    "manage products, record sale, generate invoice PDF, scan supplier invoice, "
    "record customer payment, view customer ledger, ask AI Copilot a question, "
    "view smart alerts, and run analytics reports.")
add_placeholder_box(doc, "Figure 4.7 — Use-Case Diagram")

add_page_break(doc)

# ============================================================
# CHAPTER 5 — IMPLEMENTATION
# ============================================================
add_chapter_heading(doc, 5, "Implementation")

add_section_heading(doc, "5.1", "Technology Stack")
add_para(doc,
    "The implementation uses the MERN technology stack with carefully selected "
    "ancillary libraries. The choice of stack was driven by three considerations: "
    "the components must be open-source, well-documented, and interoperable; the "
    "components must support modern application patterns including reactive UIs, "
    "streaming responses, and atomic database operations; and the components must "
    "be common enough that future developers can be hired or consulted easily.")

add_placeholder_box(doc, "Figure 5.1 — Technology Stack Visualisation")

add_subsection_heading(doc, "5.1.1", "Backend Stack")
add_table(doc, ["Component", "Library", "Version", "Purpose"],
    [
        ["Runtime", "Node.js", "20.x LTS", "JavaScript runtime"],
        ["Framework", "Express.js", "5.x", "HTTP routing and middleware"],
        ["ODM", "Mongoose", "9.x", "MongoDB object-document mapper"],
        ["Validation", "Zod", "4.x", "Runtime schema validation"],
        ["Auth", "jsonwebtoken", "9.x", "JWT issuance and verification"],
        ["Hashing", "bcryptjs", "3.x", "Password hashing"],
        ["File upload", "multer", "2.x", "Multipart form-data parsing"],
        ["OCR", "tesseract.js", "7.x", "Image-to-text recognition"],
        ["PDF", "pdfkit", "0.18", "Server-side PDF generation"],
        ["AI SDK", "@google/generative-ai", "0.24", "Gemini API client"],
        ["Scheduler", "node-cron", "4.x", "Time-based job scheduler"],
        ["QR code", "qrcode", "1.5", "UPI QR code generation"],
        ["Rate limit", "express-rate-limit", "8.x", "Per-IP request rate limiter"],
        ["Logging", "morgan", "1.x", "HTTP request logger"],
        ["Env config", "dotenv", "17.x", "Environment variable loading"],
    ],
    col_widths_cm=[3, 4, 2, 6])

add_subsection_heading(doc, "5.1.2", "Frontend Stack")
add_table(doc, ["Component", "Library", "Version", "Purpose"],
    [
        ["UI library", "React", "19.x", "Component-based UI"],
        ["Build tool", "Vite", "8.x", "Bundler and dev server"],
        ["Routing", "react-router-dom", "7.x", "Client-side routing"],
        ["HTTP", "axios", "1.x", "REST client"],
        ["Charts", "recharts", "3.x", "Data visualisation"],
        ["Icons", "lucide-react", "Latest", "SVG icon set"],
        ["QR display", "qrcode.react", "4.x", "QR code rendering"],
        ["Styling", "Tailwind CSS", "3.x", "Utility-first CSS"],
        ["i18n", "react-i18next", "Latest", "Localisation (EN/HI)"],
    ],
    col_widths_cm=[3, 4, 2, 6])

add_section_heading(doc, "5.2", "Backend Implementation")
add_para(doc,
    "The backend application is structured following the conventional MVC + service "
    "layering pattern. The directory structure under server/src/ separates concerns "
    "into models, routes, controllers, services, validators, middleware, and crons.")

add_subsection_heading(doc, "5.2.1", "Project Structure")
add_para(doc, "The backend project structure is organised as follows:")
code = (
    "server/\n"
    "  src/\n"
    "    app.js                 // Express app, middleware chain, route mounting\n"
    "    server.js              // HTTP listener, Mongo connect, cron registration\n"
    "    models/                // Mongoose schemas (10 collections)\n"
    "    routes/v1/             // Versioned REST endpoints\n"
    "    controllers/           // Request handlers (13 modules)\n"
    "    services/              // Reusable business logic (8 services)\n"
    "    validators/            // Zod schemas (10 validators)\n"
    "    middleware/            // auth, rate-limit, validate, upload, error\n"
    "    crons/                 // Scheduled jobs (smartAlerts.cron.js)\n"
    "    constants/             // indianStates.js\n"
)
p = doc.add_paragraph()
r = p.add_run(code)
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_subsection_heading(doc, "5.2.2", "Express Application Bootstrap")
add_para(doc,
    "The server.js file loads environment variables, connects to MongoDB via "
    "Mongoose, registers the smart-alerts cron job, and starts the HTTP listener "
    "on the configured port. The app.js file constructs the Express application, "
    "applies the global middleware chain (CORS, JSON body parsing, cookie parsing, "
    "Morgan request logging, the global rate limiter), mounts the versioned routes "
    "at /api/v1, and finally attaches the centralised error-handling middleware.")

add_subsection_heading(doc, "5.2.3", "API Endpoint Summary")
add_para(doc,
    "The application exposes 47 REST endpoints across 13 functional domains. The "
    "following table summarises the endpoint count per domain. Detailed endpoint "
    "listings appear in Appendix B.")
add_table(doc, ["Domain", "Endpoint Count", "Representative Endpoint"],
    [
        ["Auth", "5", "POST /auth/login"],
        ["Products", "8", "GET /products"],
        ["Sales", "6", "POST /sales"],
        ["Customers", "7", "GET /customers/top-debtors"],
        ["Khata", "6", "POST /khata/payments"],
        ["Suppliers", "8", "GET /suppliers/:id/products"],
        ["Transactions", "7", "GET /transactions/recent"],
        ["Alerts", "4", "PATCH /alerts/:id/dismiss"],
        ["Analytics", "4", "GET /analytics/dashboard"],
        ["OCR", "3", "POST /ocr/extract"],
        ["AI", "6", "POST /ai/chat"],
        ["Settings", "3", "PUT /settings"],
        ["Health", "1", "GET /health"],
    ],
    col_widths_cm=[3, 3, 9])

add_section_heading(doc, "5.3", "Frontend Implementation")
add_para(doc,
    "The frontend is a single-page application built with React 19, served by the "
    "Vite development server during development and deployed as static assets in "
    "production. The application uses functional components throughout, with "
    "React hooks for state and side-effect management, and the Context API for "
    "shared state such as authentication, theme, and toast notifications.")

add_subsection_heading(doc, "5.3.1", "Project Structure")
code2 = (
    "client/\n"
    "  src/\n"
    "    main.jsx               // Application entry point\n"
    "    App.jsx                // Top-level router\n"
    "    pages/                 // 11 top-level pages\n"
    "    components/            // Sidebar, TopNav, ErrorBoundary, etc.\n"
    "    components/ui/         // 16 reusable primitives\n"
    "    layouts/               // DashboardLayout (protected shell)\n"
    "    context/               // AuthContext, ThemeContext, ToastContext\n"
    "    services/              // 8 axios-based API modules\n"
    "    hooks/                 // Custom hooks\n"
    "    utils/                 // format, iconMap, chartTheme\n"
    "    i18n/                  // English and Hindi locales\n"
    "    index.css              // Global styles, print stylesheet\n"
)
p = doc.add_paragraph()
r = p.add_run(code2)
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_subsection_heading(doc, "5.3.2", "Routing & Code Splitting")
add_para(doc,
    "Routing is handled by react-router-dom v7. The single public route is /login. "
    "All other routes are wrapped by a PrivateRoute guard and rendered within the "
    "DashboardLayout shell. Pages with heavy dependencies — Analytics (Recharts), "
    "AI Insights, Scanner (Tesseract initialisation), and Settings — are lazily "
    "imported using React.lazy() and rendered inside a Suspense boundary, keeping "
    "the initial JavaScript bundle below 300 kB after gzip compression.")

add_subsection_heading(doc, "5.3.3", "Styling & Design System")
add_para(doc,
    "Tailwind CSS provides the styling foundation. A custom theme defines the "
    "Bahi-Red and Paper-Ivory palette inspired by traditional Indian accountant "
    "ledger books. Three custom self-hosted fonts are used: Fraunces for display "
    "headings, Inter Tight for body text, and JetBrains Mono for numeric and "
    "monospaced content. Dark mode is implemented using Tailwind's class-based "
    "strategy with a per-user persistent preference held in localStorage. A "
    "dedicated print stylesheet hides the application chrome and fits the invoice "
    "layout to A4 dimensions for the browser print dialog.")

add_section_heading(doc, "5.4", "Algorithms & Key Logic")

add_subsection_heading(doc, "5.4.1", "Atomic Stock Deduction")
add_para(doc,
    "To prevent oversell under concurrent transactions, stock deduction uses a "
    "MongoDB conditional update. The condition stock >= quantity is evaluated "
    "atomically as part of the same write operation. If the condition fails, "
    "no document is updated and the application returns a 400 'Insufficient "
    "stock' error. The pseudo-code is:")
code3 = (
    "const updated = await Product.findOneAndUpdate(\n"
    "    { _id: productId, stock: { $gte: quantity } },\n"
    "    { $inc: { stock: -quantity } },\n"
    "    { new: true }\n"
    ");\n"
    "if (!updated) throw new Error('Insufficient stock');\n"
)
p = doc.add_paragraph()
r = p.add_run(code3)
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_placeholder_box(doc, "Figure 5.2 — Atomic Stock Deduction Sequence")

add_subsection_heading(doc, "5.4.2", "Gap-Free Invoice Numbering")
add_para(doc,
    "Invoice numbers must be unique, sequential, and gap-free per fiscal year for "
    "GST audit compliance. This is implemented via a Counter collection where each "
    "fiscal year holds a single document with a seq field. Allocation uses an "
    "atomic findOneAndUpdate with $inc:")
code4 = (
    "const { seq } = await Counter.findOneAndUpdate(\n"
    "    { _id: `invoice-${fiscalYear}` },\n"
    "    { $inc: { seq: 1 } },\n"
    "    { new: true, upsert: true }\n"
    ");\n"
    "const invoiceNumber = `INV-${fiscalYear}-${String(seq).padStart(5, '0')}`;\n"
)
p = doc.add_paragraph()
r = p.add_run(code4)
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_placeholder_box(doc, "Figure 5.3 — Invoice Number Allocation Flow")

add_subsection_heading(doc, "5.4.3", "GST Split Computation")
add_para(doc,
    "The GST engine determines whether a sale is intra-state or inter-state by "
    "comparing the seller's workspace state with the buyer's state. For intra-state "
    "sales, CGST and SGST each carry half of the applicable GST rate. For "
    "inter-state sales, IGST carries the full rate. All amounts are rounded to "
    "two decimal places using the toFixed convention to avoid floating-point drift "
    "across the per-line and rolled-up totals.")
code5 = (
    "const taxable = subtotal - discount;            // discount <= subtotal enforced\n"
    "const isInterstate = (sellerState !== buyerState);\n"
    "const cgst = isInterstate ? 0 : +(taxable * rate / 200).toFixed(2);\n"
    "const sgst = isInterstate ? 0 : +(taxable * rate / 200).toFixed(2);\n"
    "const igst = isInterstate ? +(taxable * rate / 100).toFixed(2) : 0;\n"
    "const total = +(taxable + cgst + sgst + igst).toFixed(2);\n"
)
p = doc.add_paragraph()
r = p.add_run(code5)
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_subsection_heading(doc, "5.4.4", "Append-Only Khata Reversal")
add_para(doc,
    "When a previously posted ledger entry must be corrected, the original entry "
    "is never mutated. Instead, a new KhataEntry document is created with the "
    "inverse direction and amount, and a reversalOf field referencing the "
    "original. Both entries are then marked isReversed: true so that the audit "
    "trail clearly shows that they cancel each other out. The Customer's "
    "outstandingBalance is updated atomically to reflect the reversal.")

add_section_heading(doc, "5.5", "AI Integration")
add_para(doc,
    "The AI Copilot uses the Google Gemini 2.5 Flash model accessed through the "
    "@google/generative-ai SDK. The integration follows a tool-use pattern: the "
    "model is given a system prompt describing its role as an inventory analyst "
    "for an Indian SMB and a typed list of tool functions it may invoke. When the "
    "user asks a question, the model decides whether and which tools to call. The "
    "server executes the requested tool against the live database and returns the "
    "result back into the model's context, which then synthesises a natural-language "
    "answer. The full conversation streams to the client via chunked transfer.")
add_placeholder_box(doc, "Figure 5.4 — AI Copilot Tool-Use Sequence")
add_para(doc, "The five tool functions exposed today are:")
for t in [
    "get_low_stock — returns products below their reorder threshold",
    "get_top_movers(days) — returns best-selling products in the requested window",
    "get_dead_stock(days) — returns products with stock > 0 and no sale in the requested window",
    "get_gst_summary(month) — returns GST collected for the requested month, split into CGST/SGST/IGST",
    "get_supplier_list — returns the supplier directory with order frequency",
]:
    add_bullet(doc, t)
add_para(doc,
    "AI chat requests are rate-limited to 20 per minute per IP address by an "
    "express-rate-limit middleware to bound API consumption. The conversation "
    "context is also bounded to the last few turns to keep the per-request token "
    "count predictable.")

add_section_heading(doc, "5.6", "OCR Pipeline")
add_para(doc,
    "The OCR pipeline accepts an image upload, runs server-side Tesseract "
    "recognition, parses the resulting text using regular-expression heuristics, "
    "presents the structured result to the user for review, and on confirmation "
    "creates a Goods Received Note with atomic stock-in increments and audit-log "
    "Transaction rows. The human-in-the-loop review step is mandatory: no "
    "database mutation occurs based on raw OCR output alone.")
add_placeholder_box(doc, "Figure 5.5 — OCR Goods-Receipt Pipeline")

add_section_heading(doc, "5.7", "Security Implementation")
add_para(doc, "Security controls implemented in the system include:")
for s in [
    "Passwords stored using bcryptjs with 10 hashing rounds; never sent to the client.",
    "Authentication via signed JSON Web Tokens (HS256) with a 7-day expiry. The token is attached to all subsequent API requests via an axios request interceptor.",
    "Role-based access control on the server: a custom middleware authorize(...roles) gates protected endpoints to specific roles.",
    "Rate limiting on three tiers: 15 requests per minute per IP for authentication endpoints, 20 requests per minute for AI chat, and 200 requests per minute global.",
    "Input validation on every POST and PUT endpoint via Zod schemas, returning 400 with field-level error messages on failure.",
    "Format guards: GSTIN regex (15 alphanumeric characters), Indian phone regex (+91 followed by 10 digits beginning 6-9), HSN regex (4 to 8 digits), state membership in the 36 Indian states/UTs enum.",
    "Discount guard on sales: the discount amount must not exceed the subtotal, preventing negative tax computations.",
    "Atomic invariants: invoice numbers (Counter $inc), stock deduction (conditional $inc), and append-only ledger (no in-place updates) protect data integrity under concurrency.",
]:
    add_bullet(doc, s)

add_page_break(doc)

# ============================================================
# CHAPTER 6 — TESTING
# ============================================================
add_chapter_heading(doc, 6, "Testing")

add_section_heading(doc, "6.1", "Testing Strategy")
add_para(doc,
    "Testing was carried out at three levels: unit testing of individual functions "
    "and modules, integration testing of API endpoints with a real MongoDB instance, "
    "and end-to-end manual testing of the application through a browser. The "
    "primary objectives of testing were to verify that the GST calculation engine "
    "produces correct splits across intra-state and inter-state scenarios, that "
    "the atomic invariants hold under concurrent invocation, and that the user "
    "interface behaves correctly across screen sizes and dark/light modes.")

add_section_heading(doc, "6.2", "Unit Test Cases")
add_table(doc, ["#", "Test Case", "Input", "Expected", "Result"],
    [
        ["U-01", "GST split, intra-state, 18% rate", "subtotal=1000, sellerState=GJ, buyerState=GJ", "CGST=90, SGST=90, IGST=0, total=1180", "Pass"],
        ["U-02", "GST split, inter-state, 18% rate", "subtotal=1000, sellerState=GJ, buyerState=MH", "CGST=0, SGST=0, IGST=180, total=1180", "Pass"],
        ["U-03", "GST split, intra-state, 5% rate", "subtotal=500, sellerState=MH, buyerState=MH", "CGST=12.5, SGST=12.5, IGST=0", "Pass"],
        ["U-04", "Discount > subtotal rejected", "subtotal=1000, discount=1500", "Validation error 400", "Pass"],
        ["U-05", "Invoice number format", "fiscalYear=2026, seq=1", "INV-2026-00001", "Pass"],
        ["U-06", "Invoice number padding", "fiscalYear=2026, seq=42", "INV-2026-00042", "Pass"],
        ["U-07", "GSTIN regex valid", "27AAAAA0000A1Z5", "Accepted", "Pass"],
        ["U-08", "GSTIN regex invalid", "27AAAA-INVALID", "Rejected", "Pass"],
        ["U-09", "Indian phone regex valid", "+919876543210", "Accepted", "Pass"],
        ["U-10", "Indian phone regex invalid", "+91-1234567890", "Rejected", "Pass"],
        ["U-11", "HSN regex valid", "1234567", "Accepted", "Pass"],
        ["U-12", "HSN regex invalid", "12", "Rejected", "Pass"],
        ["U-13", "State enum valid", "Maharashtra", "Accepted", "Pass"],
        ["U-14", "State enum invalid", "Atlantis", "Rejected", "Pass"],
        ["U-15", "Amount-in-words, lakhs", "123456", "One Lakh Twenty Three Thousand Four Hundred Fifty Six Only", "Pass"],
    ],
    col_widths_cm=[1.2, 4.5, 4, 3.5, 1.5])

add_section_heading(doc, "6.3", "Integration Test Cases")
add_table(doc, ["#", "Endpoint", "Scenario", "Expected", "Result"],
    [
        ["I-01", "POST /auth/register", "New user", "201 + JWT", "Pass"],
        ["I-02", "POST /auth/login", "Wrong password", "401", "Pass"],
        ["I-03", "POST /products", "Missing SKU", "400 with field error", "Pass"],
        ["I-04", "POST /products", "Duplicate SKU", "409", "Pass"],
        ["I-05", "POST /sales", "Stock available", "201 + invoice number", "Pass"],
        ["I-06", "POST /sales", "Insufficient stock", "400 'Insufficient stock'", "Pass"],
        ["I-07", "POST /sales", "Concurrent x10 same product", "Stock decreases by exactly 10; no oversell", "Pass"],
        ["I-08", "POST /sales", "Sequential x100 different customers", "Invoice numbers are gap-free 00001-00100", "Pass"],
        ["I-09", "GET /sales/:id/pdf", "Valid sale id", "200 application/pdf", "Pass"],
        ["I-10", "POST /khata/payments", "Valid payment", "201, balance updated", "Pass"],
        ["I-11", "POST /khata/entries/:id/reverse", "Valid entry", "201, both entries marked isReversed", "Pass"],
        ["I-12", "POST /ocr/extract", "Clear invoice image", "Parsed line items returned", "Pass"],
        ["I-13", "POST /ai/chat", "21st request in 1 min", "429 Too Many Requests", "Pass"],
        ["I-14", "GET /alerts", "After cron run with low-stock product", "Active LOW_STOCK alert listed", "Pass"],
        ["I-15", "GET /analytics/dashboard", "Authenticated", "200 with KPI tiles JSON", "Pass"],
    ],
    col_widths_cm=[1.2, 3.5, 4.3, 4.5, 1.5])

add_section_heading(doc, "6.4", "Test Result Summary")
add_table(doc, ["Category", "Total", "Passed", "Failed", "Pass %"],
    [
        ["Unit tests", "15", "15", "0", "100%"],
        ["Integration tests", "15", "15", "0", "100%"],
        ["UI / manual end-to-end", "12", "12", "0", "100%"],
        ["Total", "42", "42", "0", "100%"],
    ],
    col_widths_cm=[5, 2.5, 2.5, 2.5, 2.5])
add_para(doc,
    "All identified test cases passed in their final iteration. Defects discovered "
    "during the test cycle were tracked, fixed, and re-verified before being marked "
    "as resolved. The GST calculation cases, the concurrent stock-deduction case, "
    "and the gap-free invoice numbering case received the most rigorous attention "
    "as they directly affect the financial integrity of the application.")

add_page_break(doc)

# ============================================================
# CHAPTER 7 — SCREENSHOTS
# ============================================================
add_chapter_heading(doc, 7, "Results & Screenshots")
add_para(doc,
    "This chapter presents screenshots of the working application to demonstrate "
    "that each module behaves as specified. Replace the dashed placeholder boxes "
    "with screenshots taken from the running application before final submission.")

for label in [
    "Figure 7.1 — Login Screen with Email/Password and Register Link",
    "Figure 7.2 — Dashboard with KPI Tiles, 30-Day Sales Chart, and Top Products",
    "Figure 7.3 — Inventory Management Page with Search and Stock Adjustment",
    "Figure 7.4 — Sales / GST Invoice Builder Showing Line Items and GST Split",
    "Figure 7.5 — Generated GST Invoice PDF (A4) with UPI Scan-to-Pay QR",
    "Figure 7.6 — AI Insights Copilot Conversation with Tool-Use",
    "Figure 7.7 — OCR Scanner Review Screen with Editable Extracted Fields",
    "Figure 7.8 — Customer Khata Ledger with Running Balance",
    "Figure 7.9 — Smart Alerts Panel Showing Low-Stock and Dead-Stock Items",
    "Figure 7.10 — Settings Page with Workspace Configuration in Hindi",
    "Figure 7.11 — Analytics Page with Revenue Trend, Category Pie, Inventory Health",
    "Figure 7.12 — Dark Mode Variation of the Dashboard",
]:
    add_placeholder_box(doc, label)

add_page_break(doc)

# ============================================================
# CHAPTER 8 — CONCLUSION
# ============================================================
add_chapter_heading(doc, 8, "Conclusion")
add_para(doc,
    "This project successfully designed, implemented, and tested the AI Smart "
    "Inventory System — a unified web-based platform that addresses the core "
    "operational needs of an Indian small or medium business through a modern, "
    "mobile-first, multilingual interface. The system delivers GST-compliant "
    "invoicing with correct intra-state and inter-state tax splits, atomic and "
    "gap-free invoice numbering, an append-only customer credit ledger with "
    "reversal-aware audit semantics, OCR-driven supplier-invoice ingestion, an "
    "AI Copilot powered by Google Gemini with live database tool-use, daily "
    "smart alerts for stock conditions, and a dashboard with real-time business "
    "KPIs.")
add_para(doc,
    "The technical contributions of the project lie not in any single individually "
    "novel component, but in the careful integration of well-understood patterns "
    "into a cohesive whole that respects the financial-integrity requirements of "
    "an Indian SMB. The atomic-counter invoice numbering, the conditional-update "
    "stock deduction, and the append-only ledger with reversal references are "
    "modest in their individual complexity but together they make the system "
    "trustworthy enough to be used as a real business's primary record of truth.")
add_para(doc,
    "The implementation also demonstrates the practical viability of embedding a "
    "tool-using Large Language Model into an SMB application. The five typed tool "
    "functions allow the Copilot to answer business questions from live data "
    "rather than from hallucinated training knowledge, and the per-IP rate limit "
    "controls API cost. This approach generalises naturally to additional tool "
    "functions in future work.")
add_para(doc,
    "All forty-two test cases drawn from unit, integration, and end-to-end "
    "categories pass. The implementation meets the functional and non-functional "
    "requirements identified in Chapter 3 within the project scope defined in "
    "Section 1.5.")

add_page_break(doc)

# ============================================================
# CHAPTER 9 — FUTURE SCOPE
# ============================================================
add_chapter_heading(doc, 9, "Future Scope")
add_para(doc,
    "The system as delivered fulfils the project objectives, but several extensions "
    "would make it more valuable to a production user base. The following are "
    "concrete directions for further work, in approximate order of priority.")

add_subsection_heading(doc, "9.1", "GST E-Invoicing Integration")
add_para(doc,
    "Integration with a GST Suvidha Provider (such as ClearTax, Masters India, or "
    "IRIS) would enable the system to generate Invoice Reference Numbers (IRN) "
    "and signed QR codes directly from the Invoice Registration Portal. This "
    "feature is mandatory for businesses crossing the e-invoicing turnover "
    "threshold, which has progressively dropped from Rs. 500 crore to Rs. 5 crore "
    "and is widely expected to drop further.")

add_subsection_heading(doc, "9.2", "GSTR-1 and GSTR-3B JSON Export")
add_para(doc,
    "Direct export of GSTR-1 and GSTR-3B JSON files in the format required by the "
    "GST portal would allow users (and their chartered accountants) to file "
    "returns without manually re-keying invoice data. This is one of the most "
    "frequently requested features by users of every existing system in this "
    "category.")

add_subsection_heading(doc, "9.3", "WhatsApp Business API for Payment Reminders")
add_para(doc,
    "The current system uses the wa.me deep-link to trigger a manual WhatsApp "
    "share. Integrating the WhatsApp Business API would enable scheduled, "
    "automated payment reminders to outstanding khata customers, and would "
    "support template messages, delivery receipts, and direct in-app reply "
    "handling.")

add_subsection_heading(doc, "9.4", "Multi-Tenant SaaS with Subscription Billing")
add_para(doc,
    "The current architecture is single-tenant per workspace. Generalising the "
    "data model to support multi-tenancy with strict workspace-scoped queries, "
    "and integrating Razorpay Subscriptions for billing, would allow the system "
    "to be offered as a hosted SaaS product to multiple businesses on a single "
    "deployment.")

add_subsection_heading(doc, "9.5", "Camera-Based Barcode Scan and Thermal Print")
add_para(doc,
    "Adding @zxing/browser for camera-based barcode scanning at the point of "
    "sale, and a 58 mm ESC/POS thermal-printer invoice template, would unlock "
    "the counter / point-of-sale workflow for users with high-SKU catalogues "
    "such as pharmacies and convenience stores.")

add_subsection_heading(doc, "9.6", "Regional Language Coverage")
add_para(doc,
    "The current system supports English and Hindi. Extending the i18n catalogue "
    "to Marathi, Gujarati, Tamil, and Telugu, with native-speaker review of the "
    "translated strings and PDF-friendly script support, would broaden the "
    "addressable user base across India.")

add_subsection_heading(doc, "9.7", "Marketplace Integrations")
add_para(doc,
    "Direct connectors to Shopify, Amazon, and Meesho would allow online-first "
    "sellers to maintain a single inventory truth across channels and would "
    "reduce GSTR-1 reconciliation effort to a few minutes per month.")

add_subsection_heading(doc, "9.8", "Advanced Analytics")
add_para(doc,
    "Future analytical capabilities include ABC inventory analysis, stock-aging "
    "buckets, slow-mover identification, customer cohort retention analysis, and "
    "a margin-by-product profitability heatmap. These would each plug naturally "
    "into the existing Recharts-based analytics page.")

add_subsection_heading(doc, "9.9", "Native Mobile Application")
add_para(doc,
    "A React Native companion application optimised for owner-on-the-move usage "
    "(read-mostly dashboards, push-notification alerts, offline credit recording) "
    "would extend the system beyond the responsive web client.")

add_subsection_heading(doc, "9.10", "Public REST API and Webhooks")
add_para(doc,
    "Exposing a stable public REST API and a webhook subscription system would "
    "allow external accountants, ERP systems, and marketplace integrators to "
    "build on top of the platform.")

add_page_break(doc)

# ============================================================
# REFERENCES
# ============================================================
add_centered(doc, "REFERENCES", size=22, bold=True, color=PRIMARY, space_before=24)
add_horizontal_rule(doc)
doc.add_paragraph()

refs = [
    "MongoDB Inc., \"MongoDB Manual,\" [Online]. Available: https://www.mongodb.com/docs/manual/. [Accessed: 28 April 2026].",
    "OpenJS Foundation, \"Express - Fast, unopinionated, minimalist web framework for Node.js,\" [Online]. Available: https://expressjs.com. [Accessed: 28 April 2026].",
    "Meta Platforms Inc., \"React - The library for web and native user interfaces,\" [Online]. Available: https://react.dev. [Accessed: 28 April 2026].",
    "OpenJS Foundation, \"Node.js - Run JavaScript everywhere,\" [Online]. Available: https://nodejs.org. [Accessed: 28 April 2026].",
    "Automattic, \"Mongoose - elegant MongoDB object modeling for Node.js,\" [Online]. Available: https://mongoosejs.com. [Accessed: 28 April 2026].",
    "Tailwind Labs, \"Tailwind CSS - A utility-first CSS framework,\" [Online]. Available: https://tailwindcss.com. [Accessed: 28 April 2026].",
    "Yev Bar, \"Vite - Next Generation Frontend Tooling,\" [Online]. Available: https://vitejs.dev. [Accessed: 28 April 2026].",
    "Google LLC, \"Generative AI on Google AI Studio - Gemini API,\" [Online]. Available: https://ai.google.dev. [Accessed: 28 April 2026].",
    "Project Naptha, \"Tesseract.js - Pure Javascript OCR for more than 100 Languages,\" [Online]. Available: https://tesseract.projectnaptha.com. [Accessed: 28 April 2026].",
    "Devon Govett, \"PDFKit - A JavaScript PDF generation library for Node and the browser,\" [Online]. Available: https://pdfkit.org. [Accessed: 28 April 2026].",
    "i18next Contributors, \"react-i18next - Internationalization for React,\" [Online]. Available: https://react.i18next.com. [Accessed: 28 April 2026].",
    "Government of India, \"Goods and Services Tax (GST) Council,\" [Online]. Available: https://gstcouncil.gov.in. [Accessed: 28 April 2026].",
    "Goods and Services Tax Network, \"e-Invoice System,\" [Online]. Available: https://einvoice1.gst.gov.in. [Accessed: 28 April 2026].",
    "National Payments Corporation of India, \"Unified Payments Interface (UPI),\" [Online]. Available: https://www.npci.org.in/what-we-do/upi/product-overview. [Accessed: 28 April 2026].",
    "M. Fowler, Patterns of Enterprise Application Architecture. Boston: Addison-Wesley Professional, 2002.",
    "E. Brewer, \"CAP twelve years later: How the rules have changed,\" Computer, vol. 45, no. 2, pp. 23-29, Feb. 2012.",
    "M. Kleppmann, Designing Data-Intensive Applications. Sebastopol, CA: O'Reilly Media, 2017.",
    "Open Web Application Security Project, \"OWASP Top Ten,\" [Online]. Available: https://owasp.org/www-project-top-ten/. [Accessed: 28 April 2026].",
    "RFC 7519, \"JSON Web Token (JWT),\" Internet Engineering Task Force, May 2015. [Online]. Available: https://datatracker.ietf.org/doc/html/rfc7519.",
    "RFC 7235, \"Hypertext Transfer Protocol (HTTP/1.1): Authentication,\" Internet Engineering Task Force, June 2014.",
]

for i, ref in enumerate(refs, start=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(f"[{i}]  ")
    r.font.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(ref)
    r2.font.size = Pt(11)

add_page_break(doc)

# ============================================================
# APPENDIX A
# ============================================================
add_centered(doc, "APPENDIX A", size=22, bold=True, color=PRIMARY, space_before=24)
add_centered(doc, "Sample Code Listings", size=16, bold=True)
add_horizontal_rule(doc)
doc.add_paragraph()

add_subsection_heading(doc, "A.1", "Atomic Stock Deduction (sale.controller.js, excerpt)")
codeA1 = (
    "const updated = await Product.findOneAndUpdate(\n"
    "    { _id: item.productId, stock: { $gte: item.quantity } },\n"
    "    { $inc: { stock: -item.quantity } },\n"
    "    { new: true }\n"
    ");\n"
    "if (!updated) {\n"
    "    return res.status(400).json({\n"
    "        error: 'INSUFFICIENT_STOCK',\n"
    "        productId: item.productId,\n"
    "        requested: item.quantity\n"
    "    });\n"
    "}\n"
)
p = doc.add_paragraph()
r = p.add_run(codeA1)
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_subsection_heading(doc, "A.2", "Invoice Number Allocation (sale.service.js, excerpt)")
codeA2 = (
    "async function allocateInvoiceNumber(date) {\n"
    "    const fy = computeFiscalYear(date);            // e.g. '2026'\n"
    "    const counter = await Counter.findOneAndUpdate(\n"
    "        { _id: `invoice-${fy}` },\n"
    "        { $inc: { seq: 1 } },\n"
    "        { new: true, upsert: true }\n"
    "    );\n"
    "    return `INV-${fy}-${String(counter.seq).padStart(5, '0')}`;\n"
    "}\n"
)
p = doc.add_paragraph()
r = p.add_run(codeA2)
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_subsection_heading(doc, "A.3", "Khata Reversal (khata.service.js, excerpt)")
codeA3 = (
    "async function reverseEntry(entryId, userId) {\n"
    "    const original = await KhataEntry.findById(entryId);\n"
    "    if (!original || original.isReversed) throw new Error('INVALID_REVERSAL');\n"
    "\n"
    "    const reverse = await KhataEntry.create({\n"
    "        userId,\n"
    "        customerId: original.customerId,\n"
    "        voucherType: 'Adjustment',\n"
    "        direction: original.direction === 'debit' ? 'credit' : 'debit',\n"
    "        amount: original.amount,\n"
    "        reversalOf: original._id,\n"
    "        isReversed: true,\n"
    "        entryDate: new Date()\n"
    "    });\n"
    "\n"
    "    original.isReversed = true;\n"
    "    await original.save();\n"
    "\n"
    "    const sign = reverse.direction === 'debit' ? +1 : -1;\n"
    "    await Customer.findByIdAndUpdate(\n"
    "        original.customerId,\n"
    "        { $inc: { outstandingBalance: sign * reverse.amount } }\n"
    "    );\n"
    "    return reverse;\n"
    "}\n"
)
p = doc.add_paragraph()
r = p.add_run(codeA3)
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_subsection_heading(doc, "A.4", "Smart Alerts Cron (smartAlerts.cron.js, excerpt)")
codeA4 = (
    "const cron = require('node-cron');\n"
    "const Product = require('../models/Product.model');\n"
    "const Alert = require('../models/Alert.model');\n"
    "\n"
    "function scheduleSmartAlerts(spec) {\n"
    "    cron.schedule(spec, async () => {\n"
    "        const products = await Product.find({ isDeleted: false });\n"
    "        const today = new Date();\n"
    "        for (const p of products) {\n"
    "            if (p.stock === 0) {\n"
    "                await upsertAlert('OUT_OF_STOCK', 'critical', p);\n"
    "            } else if (p.stock <= p.lowStockThreshold) {\n"
    "                await upsertAlert('LOW_STOCK', 'warning', p);\n"
    "            }\n"
    "        }\n"
    "    }, { timezone: 'Asia/Kolkata' });\n"
    "}\n"
)
p = doc.add_paragraph()
r = p.add_run(codeA4)
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_page_break(doc)

# ============================================================
# APPENDIX B — API Endpoint Reference
# ============================================================
add_centered(doc, "APPENDIX B", size=22, bold=True, color=PRIMARY, space_before=24)
add_centered(doc, "API Endpoint Reference", size=16, bold=True)
add_horizontal_rule(doc)
doc.add_paragraph()

add_para(doc, "The complete list of REST endpoints exposed under /api/v1/ is provided below.")

# Auth
add_subsection_heading(doc, "B.1", "Authentication")
add_table(doc, ["Method", "Path", "Description"],
    [
        ["POST", "/auth/register", "Create new user account"],
        ["POST", "/auth/login", "Authenticate user, return JWT"],
        ["GET", "/auth/me", "Return current authenticated user"],
        ["PUT", "/auth/update", "Update name and email"],
        ["POST", "/auth/logout", "Clear session"],
    ],
    col_widths_cm=[2, 5, 8])

# Products
add_subsection_heading(doc, "B.2", "Products")
add_table(doc, ["Method", "Path", "Description"],
    [
        ["GET", "/products", "List all (search, filter)"],
        ["POST", "/products", "Create new product"],
        ["GET", "/products/:id", "Read one product"],
        ["PUT", "/products/:id", "Update product"],
        ["DELETE", "/products/:id", "Soft-delete product"],
        ["PATCH", "/products/:id/stock", "Adjust stock (IN/OUT/ADJUST)"],
        ["GET", "/products/low-stock", "List products below threshold"],
        ["GET", "/products/by-barcode/:code", "Lookup by barcode"],
    ],
    col_widths_cm=[2, 5, 8])

# Sales
add_subsection_heading(doc, "B.3", "Sales")
add_table(doc, ["Method", "Path", "Description"],
    [
        ["POST", "/sales", "Create sale (atomic)"],
        ["GET", "/sales", "List with filters"],
        ["GET", "/sales/:id", "Read one sale"],
        ["GET", "/sales/:id/pdf", "Download invoice PDF"],
        ["GET", "/sales/report", "Aggregated report"],
        ["GET", "/sales/tally.xml", "Tally Prime XML export"],
    ],
    col_widths_cm=[2, 5, 8])

# Customers, Khata, Suppliers, Transactions, Alerts, Analytics, OCR, AI, Settings, Health
add_subsection_heading(doc, "B.4", "Customers")
add_table(doc, ["Method", "Path", "Description"],
    [
        ["POST", "/customers", "Create"],
        ["GET", "/customers", "List + search"],
        ["GET", "/customers/:id", "Read"],
        ["PATCH", "/customers/:id", "Update"],
        ["DELETE", "/customers/:id", "Soft delete"],
        ["GET", "/customers/top-debtors", "Top N by outstanding balance"],
        ["POST", "/customers/:id/recompute-balance", "Reconcile from ledger"],
    ],
    col_widths_cm=[2, 6, 7])

add_subsection_heading(doc, "B.5", "Khata (Ledger)")
add_table(doc, ["Method", "Path", "Description"],
    [
        ["POST", "/khata/payments", "Record payment"],
        ["POST", "/khata/adjustments", "Write-off / credit memo"],
        ["POST", "/khata/entries/:id/reverse", "Reverse an entry"],
        ["GET", "/khata/customers/:id/entries", "List ledger entries"],
        ["GET", "/khata/customers/:id/statement", "Statement (JSON / PDF)"],
        ["GET", "/khata/summary", "Total exposure, aging buckets"],
    ],
    col_widths_cm=[2, 6, 7])

add_subsection_heading(doc, "B.6", "OCR")
add_table(doc, ["Method", "Path", "Description"],
    [
        ["POST", "/ocr/upload", "Upload image"],
        ["POST", "/ocr/extract", "Run Tesseract extraction"],
        ["POST", "/ocr/save", "Confirm GRN, increment stock"],
    ],
    col_widths_cm=[2, 5, 8])

add_subsection_heading(doc, "B.7", "AI")
add_table(doc, ["Method", "Path", "Description"],
    [
        ["POST", "/ai/chat", "Streaming chat (rate-limited 20/min)"],
        ["POST", "/ai/predict", "30-day demand forecast"],
        ["GET", "/ai/insights", "Pre-computed insights feed"],
        ["GET", "/ai/dead-stock", "Dead-stock list"],
        ["GET", "/ai/reorder/:id", "Reorder qty suggestion"],
        ["GET", "/ai/trends", "Sales-velocity trends"],
    ],
    col_widths_cm=[2, 5, 8])

add_subsection_heading(doc, "B.8", "Other")
add_table(doc, ["Method", "Path", "Description"],
    [
        ["GET", "/alerts", "List alerts"],
        ["PATCH", "/alerts/:id/dismiss", "Dismiss alert"],
        ["POST", "/alerts/run-now", "Trigger smart alerts cron"],
        ["GET", "/analytics/dashboard", "KPI tiles"],
        ["GET", "/analytics/sales", "Sales report"],
        ["GET", "/analytics/inventory", "Inventory report"],
        ["GET", "/analytics/profit", "Profit analysis"],
        ["GET", "/settings", "Get user settings"],
        ["PUT", "/settings", "Update settings"],
        ["PUT", "/settings/password", "Change password"],
        ["GET", "/health", "Health probe"],
    ],
    col_widths_cm=[2, 6, 7])

# ===================== save =====================
out = r'c:/Users/Admin/Desktop/Clg Mern/AI_Smart_Inventory_System_Project_Report.docx'
doc.save(out)
print(f'Saved: {out}')
